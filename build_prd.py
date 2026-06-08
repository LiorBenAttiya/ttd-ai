from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page setup ──────────────────────────────────────────────
section = doc.sections[0]
section.page_width  = Inches(8.5)
section.page_height = Inches(11)
section.left_margin = section.right_margin = Inches(1)
section.top_margin  = section.bottom_margin = Inches(1)

# ── Colour palette ───────────────────────────────────────────
C_DARK   = RGBColor(0x1a, 0x1a, 0x2e)
C_INDIGO = RGBColor(0x63, 0x66, 0xf1)
C_CYAN   = RGBColor(0x38, 0xbd, 0xf8)
C_GREEN  = RGBColor(0x34, 0xd3, 0x99)
C_AMBER  = RGBColor(0xf5, 0x9e, 0x0b)
C_RED    = RGBColor(0xef, 0x44, 0x44)
C_GRAY   = RGBColor(0x6b, 0x72, 0x80)
C_WHITE  = RGBColor(0xff, 0xff, 0xff)

# ── Helper: set paragraph border (divider) ───────────────────
def add_divider(hex_color='6366f1'):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), hex_color)
    pBdr.append(bottom)
    pPr.append(pBdr)

def shade_cell(cell, hex_fill):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_fill)
    tcPr.append(shd)

def set_cell_border(cell, color='c7d2fe'):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top','left','bottom','right']:
        tag = OxmlElement(f'w:{side}')
        tag.set(qn('w:val'), 'single')
        tag.set(qn('w:sz'), '2')
        tag.set(qn('w:space'), '0')
        tag.set(qn('w:color'), color)
        tcBorders.append(tag)
    tcPr.append(tcBorders)

def no_border_table(table):
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for side in ['top','left','bottom','right','insideH','insideV']:
        tag = OxmlElement(f'w:{side}')
        tag.set(qn('w:val'), 'none')
        tblBorders.append(tag)
    tblPr.append(tblBorders)

def add_heading(text, level=1, color=None, page_break=False):
    if page_break:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        run = p.add_run()
        br = OxmlElement('w:br')
        br.set(qn('w:type'), 'page')
        run._r.append(br)
    h = doc.add_heading(text, level=level)
    h.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    h.paragraph_format.space_after  = Pt(4)
    if color:
        for run in h.runs:
            run.font.color.rgb = color
    return h

def add_body(text, bold=False, color=None, size=10, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(space_after)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color
    return p

def add_bullet(text, color=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.left_indent  = Inches(0.25)
    run = p.add_run(text)
    run.font.size = Pt(10)
    if color:
        run.font.color.rgb = color
    return p

def add_numbered(text, color=None):
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.left_indent  = Inches(0.25)
    run = p.add_run(text)
    run.font.size = Pt(10)
    if color:
        run.font.color.rgb = color
    return p

def add_table(rows_data, col_widths, header=None, hdr_fill='1e1b4b', border_color='c7d2fe'):
    tbl = doc.add_table(rows=0, cols=2)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    if header:
        hrow = tbl.add_row()
        hrow.cells[0].merge(hrow.cells[1])
        shade_cell(hrow.cells[0], hdr_fill)
        hp = hrow.cells[0].paragraphs[0]
        hr_ = hp.add_run(header)
        hr_.bold = True
        hr_.font.size = Pt(10)
        hr_.font.color.rgb = C_WHITE
        hp.paragraph_format.space_before = Pt(4)
        hp.paragraph_format.space_after  = Pt(4)
        hp.paragraph_format.left_indent  = Inches(0.1)
        set_cell_border(hrow.cells[0], border_color)
    for left, right in rows_data:
        row = tbl.add_row()
        lc, rc = row.cells[0], row.cells[1]
        lc.width = Inches(col_widths[0])
        rc.width = Inches(col_widths[1])
        shade_cell(lc, 'f1f5f9')
        shade_cell(rc, 'ffffff')
        set_cell_border(lc, border_color)
        set_cell_border(rc, border_color)
        lp = lc.paragraphs[0]
        lr = lp.add_run(left)
        lr.bold = True
        lr.font.size = Pt(9)
        lr.font.color.rgb = C_INDIGO
        lp.paragraph_format.space_before = Pt(4)
        lp.paragraph_format.space_after  = Pt(4)
        lp.paragraph_format.left_indent  = Inches(0.1)
        rp = rc.paragraphs[0]
        rr = rp.add_run(right)
        rr.font.size = Pt(10)
        rp.paragraph_format.space_before = Pt(4)
        rp.paragraph_format.space_after  = Pt(4)
        rp.paragraph_format.left_indent  = Inches(0.1)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


# ════════════════════════════════════════════════════════════
#  TITLE PAGE
# ════════════════════════════════════════════════════════════
def shade_paragraph(para, hex_fill):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_fill)
    pPr.append(shd)

for line_text, size, color, space_before, space_after in [
    ('PRODUCT ARCHITECTURE DOCUMENT', 22, C_WHITE,  48, 4),
    ('Personal Task Tracker — TTD',   16, C_CYAN,    0, 4),
    ('WhatsApp-First  ·  AI Agent  ·  Voice  ·  Meetings  ·  Multi-Mailbox  ·  Web', 10, C_GRAY, 0, 2),
    ('Owner: Lior   |   Version 1.0   |   June 2026   |   UI: Gemini', 9, C_GRAY, 0, 60),
]:
    p = doc.add_paragraph()
    shade_paragraph(p, '1a1a2e')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    r = p.add_run(line_text)
    r.bold = (size >= 16)
    r.font.size = Pt(size)
    r.font.color.rgb = color


# ════════════════════════════════════════════════════════════
#  1. EXECUTIVE SUMMARY
# ════════════════════════════════════════════════════════════
add_heading('1. Executive Summary', 1, C_INDIGO, page_break=True)
add_divider('6366f1')
add_body(
    'TTD (Task Tracker Dashboard) is a personal productivity system built around WhatsApp as its '
    'primary interface. The system lets the owner capture tasks via text or voice note, have an AI '
    'agent watch incoming messages and propose tasks automatically, schedule meetings from natural '
    'language, and receive a daily digest — all without opening a separate app. A web dashboard '
    '(UI delivered by Gemini) and Outlook calendar sync provide structured visibility across two '
    'business mailboxes with room to grow.'
)
add_body(
    'The architecture is API-first so the Gemini-designed UI can consume all capabilities through '
    'clean REST endpoints without tight coupling to backend logic.',
    color=C_GRAY
)
add_table([
    ('Owner',             'Lior'),
    ('Primary Interface', 'WhatsApp (personal account)'),
    ('Secondary',         'Web Dashboard (Gemini UI)  +  Outlook Calendar'),
    ('Mailboxes',         'lior@lbatech.com  ·  liorba@mepsltn.com  (+  extensible)'),
    ('AI Stack',          'OpenAI GPT-4o (NLP / agent)  ·  Whisper (speech-to-text)'),
    ('Messaging',         'WhatsApp Business Cloud API  /  Twilio (fallback)'),
    ('Backend',           'Python FastAPI  +  PostgreSQL (Supabase)'),
    ('Hosting',           'Railway (API)  ·  Vercel (Web)  ·  Supabase (DB + Auth)'),
    ('UI Delivery',       'Gemini — REST API consumer only, no backend involvement'),
    ('Document Version',  '1.0  —  June 2026'),
], col_widths=(2.0, 5.3), header='Project Overview')


# ════════════════════════════════════════════════════════════
#  2. SYSTEM ARCHITECTURE
# ════════════════════════════════════════════════════════════
add_heading('2. System Architecture Overview', 1, C_INDIGO, page_break=True)
add_divider('6366f1')
add_body('The system is divided into five discrete layers. Each layer communicates only with adjacent layers through defined interfaces.')

add_heading('2.1  Layer Map', 2, C_CYAN)
add_table([
    ('L1  Ingestion',     'Receives all inbound events: WhatsApp webhooks (text + audio), Outlook Graph API webhooks, web API calls from the Gemini UI.'),
    ('L2  AI Processing', 'Transcribes voice (Whisper), parses natural language (GPT-4o), classifies intent (new task / meeting / completion / query), extracts structured fields.'),
    ('L3  Domain',        'Core business logic: task CRUD, meeting scheduling, contact matching, mailbox routing, daily report generation, approval-flow state machine.'),
    ('L4  Storage',       'PostgreSQL via Supabase: tasks, mailboxes, contacts, companies, meetings, approval_queue, scheduled_reports, audit_log.'),
    ('L5  Delivery',      'Sends outbound WhatsApp messages, creates Outlook calendar invites, exposes REST API for the Gemini web UI.'),
], col_widths=(2.0, 5.3), header='Architecture Layers')

add_heading('2.2  Request Flow — Voice Note to Task', 2, C_CYAN)
for s in [
    'User sends a voice note on WhatsApp.',
    'WhatsApp Cloud API POSTs audio webhook to L1  →  /webhook/whatsapp.',
    'L2 downloads audio file, calls Whisper API  →  receives transcript text.',
    'L2 sends transcript to GPT-4o with structured JSON output prompt  →  {description, due_date, priority, source}.',
    'L3 validates fields, writes row to tasks table (status = pending).',
    'L5 sends WhatsApp confirmation:  "Got it  ✅  —  [description]  ·  Due [date]  ·  🔴 Priority 1".',
    'Task is immediately visible in web dashboard via GET /api/v1/tasks.',
]:
    add_numbered(s)
doc.add_paragraph()

add_heading('2.3  Request Flow — AI Agent Proposes Task', 2, C_CYAN)
for s in [
    'Every inbound WhatsApp message triggers L1 webhook (not only explicit commands).',
    'L2 AI Agent runs parallel intent classification: is this message actionable?',
    'If actionable: GPT-4o extracts proposed task fields + confidence score.',
    'If confidence >= 0.75: L3 writes row to approval_queue (status = proposed).',
    'L5 sends WhatsApp:  "I noticed something actionable — [proposed task]. Reply ✅ to add, ✏️ [edit] to change, ❌ to skip."',
    'User reply updates approval_queue:  approved  →  promoted to tasks table;  rejected  →  archived.',
]:
    add_numbered(s)
doc.add_paragraph()


# ════════════════════════════════════════════════════════════
#  3. DATA MODEL
# ════════════════════════════════════════════════════════════
add_heading('3. Data Model', 1, C_INDIGO, page_break=True)
add_divider('6366f1')
add_body('All tables live in PostgreSQL (Supabase). Primary keys are UUIDs. All timestamps are UTC. Tasks use soft deletes (deleted_at nullable).')

add_heading('3.1  tasks', 2, C_CYAN)
add_table([
    ('id',           'UUID  PRIMARY KEY  DEFAULT gen_random_uuid()'),
    ('description',  'TEXT  NOT NULL'),
    ('due_date',     'DATE'),
    ('priority',     'SMALLINT  NOT NULL  CHECK (priority IN (1,2,3))  — 1=Business  2=Personal  3=General'),
    ('status',       "TEXT  NOT NULL  DEFAULT 'pending'  CHECK (status IN ('pending','in_progress','completed','archived'))"),
    ('source',       "TEXT  — 'whatsapp_text' | 'whatsapp_voice' | 'web' | 'outlook_flag' | 'ai_agent'"),
    ('mailbox_id',   'UUID  FK → mailboxes.id  NULLABLE'),
    ('contact_id',   'UUID  FK → contacts.id   NULLABLE'),
    ('meeting_id',   'UUID  FK → meetings.id   NULLABLE'),
    ('created_at',   'TIMESTAMPTZ  DEFAULT now()'),
    ('completed_at', 'TIMESTAMPTZ  NULLABLE'),
    ('deleted_at',   'TIMESTAMPTZ  NULLABLE  — soft delete'),
], col_widths=(2.0, 5.3), header='Table: tasks')

add_heading('3.2  mailboxes', 2, C_CYAN)
add_table([
    ('id',                  'UUID  PRIMARY KEY'),
    ('email',               'TEXT  UNIQUE  NOT NULL'),
    ('display_name',        'TEXT  NOT NULL'),
    ('oauth_access_token',  'TEXT  ENCRYPTED (Supabase Vault)'),
    ('oauth_refresh_token', 'TEXT  ENCRYPTED (Supabase Vault)'),
    ('token_expires_at',    'TIMESTAMPTZ'),
    ('active',              'BOOLEAN  DEFAULT TRUE'),
    ('created_at',          'TIMESTAMPTZ  DEFAULT now()'),
], col_widths=(2.0, 5.3), header='Table: mailboxes')

add_heading('3.3  companies  &  contacts', 2, C_CYAN)
add_table([
    ('companies.id',       'UUID  PRIMARY KEY'),
    ('companies.name',     'TEXT  NOT NULL'),
    ('companies.domain',   'TEXT'),
    ('contacts.id',        'UUID  PRIMARY KEY'),
    ('contacts.company_id','UUID  FK → companies.id'),
    ('contacts.full_name', 'TEXT  NOT NULL'),
    ('contacts.email',     'TEXT'),
    ('contacts.phone',     'TEXT'),
    ('contacts.whatsapp',  'TEXT'),
    ('contacts.role',      'TEXT'),
], col_widths=(2.0, 5.3), header='Tables: companies & contacts')

add_heading('3.4  meetings', 2, C_CYAN)
add_table([
    ('id',               'UUID  PRIMARY KEY'),
    ('title',            'TEXT  NOT NULL'),
    ('start_time',       'TIMESTAMPTZ  NOT NULL'),
    ('end_time',         'TIMESTAMPTZ'),
    ('outlook_event_id', 'TEXT  — MS Graph event ID for updates & cancellations'),
    ('mailbox_id',       'UUID  FK → mailboxes.id  — which account sends the invite'),
    ('attendees',        'JSONB  — [{email, name, response_status}]'),
    ('status',           "TEXT  DEFAULT 'pending'  CHECK (status IN ('pending','sent','confirmed','cancelled'))"),
    ('task_id',          'UUID  FK → tasks.id  NULLABLE'),
], col_widths=(2.0, 5.3), header='Table: meetings')

add_heading('3.5  approval_queue', 2, C_CYAN)
add_table([
    ('id',               'UUID  PRIMARY KEY'),
    ('proposed_task',    'JSONB  — {description, due_date, priority, source_message_text}'),
    ('confidence',       'FLOAT  — GPT-4o confidence 0.0–1.0'),
    ('wa_message_id',    'TEXT   — original WhatsApp message ID'),
    ('status',           "TEXT  DEFAULT 'proposed'  CHECK (status IN ('proposed','approved','edited','rejected','expired'))"),
    ('resolved_task_id', 'UUID  FK → tasks.id  NULLABLE  — set on approval'),
    ('created_at',       'TIMESTAMPTZ  DEFAULT now()'),
    ('resolved_at',      'TIMESTAMPTZ  NULLABLE'),
], col_widths=(2.0, 5.3), header='Table: approval_queue')


# ════════════════════════════════════════════════════════════
#  4. REST API
# ════════════════════════════════════════════════════════════
add_heading('4. REST API Design', 1, C_INDIGO, page_break=True)
add_divider('6366f1')
add_body('Base path: /api/v1. Auth: Bearer JWT (Supabase Auth). Webhooks use /webhook namespace with HMAC verification. OpenAPI spec auto-generated at /docs.')

add_heading('4.1  Webhooks', 2, C_CYAN)
add_table([
    ('POST /webhook/whatsapp', 'Receive WA events (text, audio, status updates). Verify HMAC. Dispatch to L2 processor.'),
    ('GET  /webhook/whatsapp', 'WhatsApp hub.challenge verification on first setup.'),
    ('POST /webhook/outlook',  'Receive MS Graph change notifications (mail flag changes).'),
], col_widths=(2.8, 4.5), header='Webhook Endpoints')

add_heading('4.2  Tasks', 2, C_CYAN)
add_table([
    ('GET    /tasks',              'List tasks. Params: status, priority, due_before, due_after, mailbox_id, limit, offset.'),
    ('POST   /tasks',              'Create task. Body: {description, due_date, priority, mailbox_id?, contact_id?}'),
    ('GET    /tasks/{id}',         'Single task with linked contact, meeting, mailbox.'),
    ('PATCH  /tasks/{id}',         'Partial update of any task field.'),
    ('DELETE /tasks/{id}',         'Soft-delete (sets deleted_at).'),
    ('POST   /tasks/{id}/complete','Set status=completed, completed_at=now().'),
], col_widths=(2.8, 4.5), header='Task Endpoints')

add_heading('4.3  Approval Queue', 2, C_CYAN)
add_table([
    ('GET   /approvals',             'List proposed tasks where status=proposed.'),
    ('POST  /approvals/{id}/approve','Approve. Accepts optional override fields. Creates tasks row.'),
    ('POST  /approvals/{id}/reject', 'Reject. Sets status=rejected, no task created.'),
    ('PATCH /approvals/{id}/edit',   'Edit proposed fields before final approval.'),
], col_widths=(2.8, 4.5), header='Approval Queue Endpoints')

add_heading('4.4  Meetings', 2, C_CYAN)
add_table([
    ('POST   /meetings',      'Create meeting + send Outlook invites. Body: {title, start_time, end_time, attendees[], mailbox_id}.'),
    ('GET    /meetings',      'List meetings. Params: status, from, to.'),
    ('PATCH  /meetings/{id}', 'Update meeting (syncs Outlook event via Graph API).'),
    ('DELETE /meetings/{id}', 'Cancel meeting. Sends Outlook cancellation. status=cancelled.'),
], col_widths=(2.8, 4.5), header='Meeting Endpoints')

add_heading('4.5  Mailboxes, Contacts & Reports', 2, C_CYAN)
add_table([
    ('GET    /mailboxes',              'List connected mailboxes.'),
    ('POST   /mailboxes/connect',      'Start OAuth flow to connect a new Outlook mailbox.'),
    ('DELETE /mailboxes/{id}',         'Disconnect mailbox.'),
    ('GET    /companies',              'List companies with contact counts.'),
    ('POST   /companies/{id}/import',  'Upload CSV or vCard — bulk import contacts for a company.'),
    ('GET    /contacts',               'Search contacts by name, email, or company_id.'),
    ('GET    /reports/daily',          'Preview next daily report as JSON.'),
    ('POST   /reports/send-now',       'Trigger immediate WhatsApp report.'),
    ('PATCH  /reports/schedule',       'Update cron expression, timezone, look_ahead days.'),
], col_widths=(2.8, 4.5), header='Mailboxes / Contacts / Reports')


# ════════════════════════════════════════════════════════════
#  5. INTEGRATIONS
# ════════════════════════════════════════════════════════════
add_heading('5. External Integrations', 1, C_INDIGO, page_break=True)
add_divider('6366f1')

add_heading('5.1  WhatsApp Business Cloud API (Meta)', 2, C_CYAN)
add_table([
    ('Inbound events',   'text, audio, document, reaction — all delivered as webhook POST'),
    ('Outbound',         'POST https://graph.facebook.com/v19.0/{phone_id}/messages'),
    ('Audio pipeline',   'Media ID  →  GET /media/{id}  →  binary  →  Whisper API'),
    ('Auth',             'Bearer token from Meta Business Manager'),
    ('Phone scope',      'System only processes messages from owner\'s registered number. All others dropped.'),
    ('Fallback',         'Twilio WhatsApp on Meta 5xx errors'),
], col_widths=(2.0, 5.3), header='WhatsApp Integration')

add_heading('5.2  OpenAI (GPT-4o + Whisper)', 2, C_CYAN)
add_table([
    ('Whisper STT',       'POST /v1/audio/transcriptions — mp3/ogg input. Language: auto-detect (Hebrew + English).'),
    ('GPT-4o Task Parse', 'POST /v1/chat/completions — system prompt enforces JSON schema: {description, due_date, priority, confidence, intent}.'),
    ('GPT-4o AI Agent',   'Separate system prompt: classify actionability, return proposed_task JSON or null.'),
    ('Meeting Parser',    'Dedicated prompt: extract {title, datetime, duration, attendees[]} from free text.'),
    ('Retry policy',      'Exponential backoff: 3 attempts, 1s / 2s / 4s. Failures logged to audit_log.'),
], col_widths=(2.0, 5.3), header='OpenAI Integration')

add_heading('5.3  Microsoft Graph API (Outlook)', 2, C_CYAN)
add_table([
    ('Auth flow',         'OAuth 2.0 Authorization Code + offline_access. Tokens encrypted in Supabase Vault.'),
    ('Token refresh',     'Background worker checks expiry every 30 min. Refreshes silently.'),
    ('Create event',      'POST /me/events — title, time, attendees, location.'),
    ('Update event',      'PATCH /me/events/{id}'),
    ('Cancel event',      'DELETE /me/events/{id} — triggers Outlook cancellation email to attendees.'),
    ('Mail flag watch',   'POST /subscriptions — webhook on mail flag change  →  auto task creation.'),
    ('Multi-mailbox',     'Each mailbox has independent OAuth token. API calls use token matching mailbox_id.'),
    ('Required scopes',   'Calendars.ReadWrite  ·  Mail.ReadWrite  ·  offline_access  ·  User.Read'),
], col_widths=(2.0, 5.3), header='MS Graph Integration')


# ════════════════════════════════════════════════════════════
#  6. AI AGENT
# ════════════════════════════════════════════════════════════
add_heading('6. AI Agent Design', 1, C_INDIGO, page_break=True)
add_divider('6366f1')
add_body('The AI Agent runs on every inbound WhatsApp message that is not an explicit command. Nothing is ever added to tasks without the owner\'s explicit approval.')

add_heading('6.1  Agent Prompt Architecture', 2, C_CYAN)
add_table([
    ('System role',         'Personal assistant reading Lior\'s WhatsApp messages. Identify commitments, action items, deadlines. Return structured JSON or null.'),
    ('Output schema',       '{ "actionable": bool, "confidence": float, "proposed": { "description": str, "due_date": "YYYY-MM-DD"|null, "priority": 1|2|3, "reasoning": str } | null }'),
    ('Confidence threshold','0.75 — below this, agent stays silent. No proposal sent to owner.'),
    ('Context window',      'Last 5 messages in the conversation thread included for context.'),
    ('Language',            'Hebrew and English supported natively by GPT-4o.'),
    ('Rate limiting',       'Max 1 proposal per conversation thread per 10 minutes.'),
], col_widths=(2.2, 5.1), header='Agent Configuration')

add_heading('6.2  Approval State Machine', 2, C_CYAN)
add_table([
    ('proposed → approved', 'User replies ✅ or "yes". Task written to tasks table. approval_queue.status = approved.'),
    ('proposed → edited',   'User replies ✏️ + new text. System re-proposes with edits for final confirmation.'),
    ('proposed → rejected', 'User replies ❌ or "no". Row archived. No task created.'),
    ('proposed → expired',  'No response within 24 hours. Auto-archived. No follow-up sent.'),
], col_widths=(2.5, 4.8), header='Approval States')

add_heading('6.3  Meeting Parser', 2, C_CYAN)
add_body('Triggered when message contains scheduling intent keywords. Dedicated GPT-4o prompt extracts:')
for item in [
    'title — inferred from message topic',
    'start_time — parsed from relative ("next Tuesday 3pm") or absolute expressions',
    'duration — defaults to 60 min if not stated',
    'attendees — fuzzy-matched against contacts table by name',
    'mailbox_id — inferred from company context or defaults to lbatech',
]:
    add_bullet(item)
add_body('Extracted meeting is sent to owner for approval via WhatsApp before any Outlook invite is dispatched.', color=C_GRAY)
doc.add_paragraph()


# ════════════════════════════════════════════════════════════
#  7. DAILY REPORT
# ════════════════════════════════════════════════════════════
add_heading('7. Daily Report Specification', 1, C_INDIGO, page_break=True)
add_divider('6366f1')

add_table([
    ('Default send time', '07:00 daily — configurable in scheduled_reports table'),
    ('Timezone',          'Asia/Jerusalem (configurable)'),
    ('Trigger',           'Cron job — Supabase Edge Functions or Railway background worker'),
    ('Delivery channel',  'WhatsApp message to owner\'s personal number'),
    ('Look-ahead window', '7 days (configurable)'),
], col_widths=(2.0, 5.3), header='Schedule Config')

add_heading('7.1  Report Message Structure', 2, C_CYAN)
add_table([
    ('Header',          'Good morning Lior!  —  [Day, Date]'),
    ('Section 1: TODAY','All tasks due today. Sorted by priority, then due time. Prefixed by emoji.'),
    ('Section 2: WEEK', 'Tasks due in next 7 days. Grouped by day. Prefixed by emoji.'),
    ('Section 3: AI',   'Count of AI-proposed tasks awaiting review.'),
    ('Section 4: MEETINGS','Upcoming calendar events this week.'),
    ('Footer',          'Yesterday: N tasks completed. Motivational note if N > 0.'),
    ('Mailbox labels',  'Each task shows source mailbox: [lbatech] or [mepsltn].'),
], col_widths=(2.0, 5.3), header='Report Structure')

add_heading('7.2  Priority Emoji Encoding', 2, C_CYAN)
for item in [
    '🔴  Priority 1 — Business',
    '🟡  Priority 2 — Personal',
    '🟢  Priority 3 — General',
    '📅  Meeting / Calendar event',
    '🤖  AI-proposed task awaiting approval',
]:
    add_bullet(item)
doc.add_paragraph()


# ════════════════════════════════════════════════════════════
#  8. CONTACT DIRECTORY
# ════════════════════════════════════════════════════════════
add_heading('8. Contact Directory', 1, C_INDIGO, page_break=True)
add_divider('6366f1')
add_body('Contacts are organised per company. Used by the AI agent for name matching and by the meeting scheduler for Outlook invite addressing.')

add_heading('8.1  Import Formats', 2, C_CYAN)
add_table([
    ('CSV',    'Columns: full_name, email, phone, whatsapp, role, notes. Header row required. Max 500 rows.'),
    ('vCard',  'vCard 3.0 / 4.0. Multiple contacts per .vcf file supported.'),
    ('Manual', 'POST /api/v1/contacts or WA command: "Add contact: Name, email, Company"'),
], col_widths=(1.5, 5.8), header='Import Methods')

add_heading('8.2  AI Name Matching Algorithm', 2, C_CYAN)
for s in [
    'Exact match on full_name (case-insensitive).',
    'Fuzzy match (Levenshtein distance <= 2) on full_name and first name.',
    'Match confidence >= 0.8: auto-link task/meeting to contact_id.',
    'Match confidence 0.5–0.79: propose match to owner via WhatsApp.',
    'Match confidence < 0.5: leave contact_id null, note in task description.',
]:
    add_numbered(s)
doc.add_paragraph()


# ════════════════════════════════════════════════════════════
#  9. MULTI-MAILBOX
# ════════════════════════════════════════════════════════════
add_heading('9. Multi-Mailbox Design', 1, C_INDIGO, page_break=True)
add_divider('6366f1')
add_body('Each Outlook mailbox is an independent OAuth identity. The system supports unlimited mailboxes. Adding one requires only an OAuth login — no code changes.')

add_table([
    ('lior@lbatech.com',    'LBATech — primary business mailbox. Default when company context is absent.'),
    ('liorba@mepsltn.com',  'Mepsltn — secondary business mailbox.'),
    ('Adding more',         'Settings > Mailboxes > Connect New. Complete Microsoft OAuth. Active immediately.'),
    ('Task routing',        'Flagged-email tasks inherit mailbox_id of source mailbox. AI tasks use contact company match.'),
    ('Meeting invites',     'Owner or AI specifies sending mailbox. Invite "from" address matches that account.'),
    ('Report display',      'Daily report tags each task with mailbox shortname for quick visual identification.'),
    ('Token management',    'Each token refreshed independently. Failure on one does not affect others.'),
], col_widths=(2.2, 5.1), header='Mailbox Configuration')


# ════════════════════════════════════════════════════════════
#  10. SECURITY
# ════════════════════════════════════════════════════════════
add_heading('10. Security Design', 1, C_INDIGO, page_break=True)
add_divider('6366f1')

add_table([
    ('Authentication',      'Supabase Auth — JWT Bearer tokens. Web login via email + OTP. No passwords stored.'),
    ('Webhook HMAC',        'Every WhatsApp webhook verified with HMAC-SHA256. Requests without valid sig return 401.'),
    ('OAuth token storage', 'Outlook tokens AES-256 encrypted at rest in Supabase Vault. Never returned to client.'),
    ('API access control',  'All /api/v1/* requires valid JWT. Single-user system — no multi-tenant risk.'),
    ('Secrets management',  'API keys (OpenAI, Meta, Twilio) stored in Railway / Vercel env vars only. Not in DB or code.'),
    ('WA number whitelist', 'System processes messages only from owner\'s registered number. All others silently dropped.'),
    ('Audit log',           'All task mutations and approvals written to audit_log with timestamp, action, old/new values.'),
    ('Transport',           'TLS 1.3 on all endpoints. HSTS enforced. No HTTP.'),
], col_widths=(2.0, 5.3), header='Security Controls')


# ════════════════════════════════════════════════════════════
#  11. BUILD PHASES
# ════════════════════════════════════════════════════════════
add_heading('11. Build Phases', 1, C_INDIGO, page_break=True)
add_divider('6366f1')

add_heading('Phase 1 — MVP: WhatsApp Core', 2, C_CYAN)
add_body('Deliverable: Working WhatsApp bot that creates tasks (text + voice) and sends a daily report.', bold=True)
for item in [
    'FastAPI project scaffold + Supabase DB setup (tasks, mailboxes, scheduled_reports)',
    'WhatsApp webhook endpoint with HMAC verification',
    'Text command parser: create task, list tasks, mark done',
    'Voice note pipeline: download audio → Whisper → GPT-4o → task row',
    'WhatsApp confirmation message on every task creation',
    'Daily report cron: query DB → format → send WhatsApp',
    'Seed with two mailboxes (lbatech + mepsltn)',
]:
    add_bullet(item)
doc.add_paragraph()

add_heading('Phase 2 — REST API for Gemini UI', 2, C_CYAN)
add_body('Deliverable: Full REST API with OpenAPI spec, ready for Gemini UI consumption.', bold=True)
for item in [
    'All /api/v1/* endpoints implemented with pagination and filtering',
    'Supabase Auth + JWT middleware on all routes',
    'CORS configured for Gemini UI domain',
    'OpenAPI / Swagger docs at /docs',
    'API key management endpoint for future integrations',
]:
    add_bullet(item)
doc.add_paragraph()

add_heading('Phase 3 — AI Agent + Meeting Scheduler', 2, C_CYAN)
add_body('Deliverable: Agent monitors all WA messages. Meetings schedulable via voice or text.', bold=True)
for item in [
    'AI agent intent classifier on every inbound WA message',
    'approval_queue table + full approval state machine',
    'WA approval flow: ✅  ✏️  ❌ reply handling',
    'Meeting parser prompt + meetings table',
    'MS Graph OAuth flow + calendar event creation + cancellation',
    'Multi-mailbox token refresh background worker',
]:
    add_bullet(item)
doc.add_paragraph()

add_heading('Phase 4 — Contact Directory + Outlook Flag Sync', 2, C_CYAN)
add_body('Deliverable: Full company/contact directory. Flagged emails auto-create tasks.', bold=True)
for item in [
    'companies + contacts tables + CSV / vCard import endpoint',
    'Fuzzy name matcher used by AI agent and meeting parser',
    'MS Graph subscription for Outlook mail flag  →  task creation',
    'Contact-linked task display in all API responses',
    'Add-mailbox OAuth self-service flow via web UI',
]:
    add_bullet(item)
doc.add_paragraph()


# ════════════════════════════════════════════════════════════
#  12. OPEN QUESTIONS
# ════════════════════════════════════════════════════════════
add_heading('12. Open Questions', 1, C_INDIGO, page_break=True)
add_divider('6366f1')
add_body('All marked [Owner] require Lior\'s decision before or during the relevant phase.')

add_table([
    ('[P1 — Owner]  WhatsApp number type',
     'Use personal WA via Twilio Business API, or apply for a dedicated Meta Cloud API number? Personal has lower throughput limits.'),
    ('[P1 — Owner]  Daily report time',
     'Confirm send time. Default is 07:00 Asia/Jerusalem.'),
    ('[P1 — Owner]  Bot language',
     'Should the bot reply in Hebrew, English, or auto-mirror the language of the incoming message?'),
    ('[P3 — Owner]  Default meeting duration',
     'When duration is not stated: 30 min or 60 min?'),
    ('[P3 — Owner]  Default invite mailbox',
     'When AI auto-creates a meeting: send invite from lbatech or mepsltn by default?'),
    ('[P4 — Owner]  Contact sample',
     'Provide a sample CSV from existing contacts so import parser can be validated against real data.'),
    ('[P2 — Owner]  Web UI domain',
     'Confirm domain for the Gemini web dashboard. Needed for CORS and TLS configuration.'),
], col_widths=(2.8, 4.5), header='Open Questions')


# ════════════════════════════════════════════════════════════
#  13. GLOSSARY
# ════════════════════════════════════════════════════════════
add_heading('13. Glossary', 1, C_INDIGO)
add_divider('6366f1')
add_table([
    ('TTD',            'Task Tracker Dashboard — the product name'),
    ('L1–L5',          'Architecture layers: Ingestion, AI Processing, Domain, Storage, Delivery'),
    ('AI Agent',       'Background process reading inbound WA messages and proposing tasks'),
    ('Approval Queue', 'Staging table for AI-proposed tasks awaiting owner confirmation'),
    ('Whisper',        'OpenAI speech-to-text model used to transcribe WhatsApp voice notes'),
    ('GPT-4o',         'OpenAI language model for NLP, intent classification, and meeting parsing'),
    ('Graph API',      'Microsoft Graph REST API for Outlook calendar events and mail reading'),
    ('Mailbox',        'A connected Outlook email account with its own OAuth token'),
    ('P1 / P2 / P3',   'Task priority: 1=Business (red), 2=Personal (yellow), 3=General (green)'),
    ('WA',             'WhatsApp'),
], col_widths=(1.5, 5.8), header='Glossary')


# ── Save ─────────────────────────────────────────────────────
out = '/sessions/brave-kind-johnson/mnt/TTD/TTD_Product_Architecture_v1.0.docx'
doc.save(out)
print(f'Saved: {out}')
