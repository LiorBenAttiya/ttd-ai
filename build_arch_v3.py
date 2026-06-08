from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

s = doc.sections[0]
s.page_width = Inches(8.5); s.page_height = Inches(11)
s.left_margin = s.right_margin = Inches(1)
s.top_margin  = s.bottom_margin = Inches(1)

C_INDIGO = RGBColor(0x63,0x66,0xf1)
C_CYAN   = RGBColor(0x38,0xbd,0xf8)
C_GREEN  = RGBColor(0x34,0xd3,0x99)
C_AMBER  = RGBColor(0xf5,0x9e,0x0b)
C_PINK   = RGBColor(0xc0,0x26,0xd3)
C_RED    = RGBColor(0xef,0x44,0x44)
C_GRAY   = RGBColor(0x6b,0x72,0x80)
C_WHITE  = RGBColor(0xff,0xff,0xff)
C_AZURE  = RGBColor(0x00,0x78,0xd4)
C_SAP    = RGBColor(0x00,0x9a,0x44)

def shade_para(p, hex_fill):
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto'); shd.set(qn('w:fill'),hex_fill)
    pPr.append(shd)

def shade_cell(cell, hex_fill):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto'); shd.set(qn('w:fill'),hex_fill)
    tcPr.append(shd)

def cell_border(cell, color='c7d2fe'):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    b = OxmlElement('w:tcBorders')
    for side in ['top','left','bottom','right']:
        t = OxmlElement(f'w:{side}')
        t.set(qn('w:val'),'single'); t.set(qn('w:sz'),'2')
        t.set(qn('w:space'),'0');    t.set(qn('w:color'),color)
        b.append(t)
    tcPr.append(b)

def divider(color='6366f1'):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'),'single'); bot.set(qn('w:sz'),'6')
    bot.set(qn('w:space'),'1');    bot.set(qn('w:color'),color)
    pBdr.append(bot); pPr.append(pBdr)

def H(text, level=1, color=None, pb=False):
    if pb:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0)
        br = OxmlElement('w:br'); br.set(qn('w:type'),'page')
        p.add_run()._r.append(br)
    h = doc.add_heading(text, level=level)
    h.paragraph_format.space_before = Pt(12 if level==1 else 8)
    h.paragraph_format.space_after  = Pt(4)
    if color:
        for r in h.runs: r.font.color.rgb = color
    return h

def body(text, bold=False, color=None, size=10, after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(after)
    r = p.add_run(text); r.font.size = Pt(size); r.bold = bold
    if color: r.font.color.rgb = color
    return p

def bullet(text, color=None, indent=0.25):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(1); p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Inches(indent)
    r = p.add_run(text); r.font.size = Pt(10)
    if color: r.font.color.rgb = color

def numbered(text):
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.space_before = Pt(1); p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Inches(0.25)
    p.add_run(text).font.size = Pt(10)

def tbl(rows, col_w=(2.0,5.3), header=None, hfill='1e1b4b', bc='c7d2fe'):
    t = doc.add_table(rows=0, cols=2); t.style='Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    if header:
        hr = t.add_row(); hr.cells[0].merge(hr.cells[1])
        shade_cell(hr.cells[0], hfill); cell_border(hr.cells[0], bc)
        hp = hr.cells[0].paragraphs[0]
        r = hp.add_run(header); r.bold=True; r.font.size=Pt(10); r.font.color.rgb=C_WHITE
        hp.paragraph_format.space_before=Pt(4); hp.paragraph_format.space_after=Pt(4)
        hp.paragraph_format.left_indent=Inches(0.1)
    for left, right in rows:
        row = t.add_row(); lc,rc = row.cells[0],row.cells[1]
        lc.width=Inches(col_w[0]); rc.width=Inches(col_w[1])
        shade_cell(lc,'f1f5f9'); shade_cell(rc,'ffffff')
        cell_border(lc,bc); cell_border(rc,bc)
        lp=lc.paragraphs[0]; lr=lp.add_run(left)
        lr.bold=True; lr.font.size=Pt(9); lr.font.color.rgb=C_INDIGO
        lp.paragraph_format.space_before=Pt(4); lp.paragraph_format.space_after=Pt(4)
        lp.paragraph_format.left_indent=Inches(0.1)
        rp=rc.paragraphs[0]; rr=rp.add_run(right); rr.font.size=Pt(10)
        rp.paragraph_format.space_before=Pt(4); rp.paragraph_format.space_after=Pt(4)
        rp.paragraph_format.left_indent=Inches(0.1)
    doc.add_paragraph().paragraph_format.space_after=Pt(4)

def sp(): doc.add_paragraph().paragraph_format.space_after=Pt(4)

def note_box(text, fill='e0f2fe', border='0078d4', color=None):
    tbl_n = doc.add_table(rows=1, cols=1); tbl_n.style='Table Grid'
    c = tbl_n.rows[0].cells[0]
    shade_cell(c, fill); cell_border(c, border)
    c.width = Inches(7.3)
    p = c.paragraphs[0]; r = p.add_run(text); r.font.size=Pt(10)
    if color: r.font.color.rgb = color
    p.paragraph_format.space_before=Pt(6); p.paragraph_format.space_after=Pt(6)
    p.paragraph_format.left_indent=Inches(0.1)
    doc.add_paragraph().paragraph_format.space_after=Pt(4)

# ════════════════════════════════════════════════════════════
#  TITLE PAGE
# ════════════════════════════════════════════════════════════
for text, size, color, sb, sa in [
    ('PRODUCT ARCHITECTURE DOCUMENT', 22, C_WHITE,  48, 4),
    ('TTD AI — Personal & Business Task Tracker',   16, C_CYAN,    0, 4),
    ('WhatsApp-First  ·  AI Agent  ·  Voice  ·  Meetings  ·  Multi-Mailbox  ·  SAP B1  ·  Azure', 10, C_GRAY, 0, 2),
    ('Version 3.0  —  June 2026  —  Hosting  ·  System Requirements  ·  Future Connectors', 9, C_AMBER, 0, 2),
    ('Owner: Lior   |   UI: Gemini   |   Infrastructure: Azure + GitHub', 9, C_GRAY, 0, 60),
]:
    p = doc.add_paragraph(); shade_para(p,'1a1a2e')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before=Pt(sb); p.paragraph_format.space_after=Pt(sa)
    r = p.add_run(text); r.bold=(size>=16); r.font.size=Pt(size); r.font.color.rgb=color


# ════════════════════════════════════════════════════════════
#  CHANGE LOG v2 → v3
# ════════════════════════════════════════════════════════════
H('Change Log  —  v2.0 → v3.0', 1, C_AMBER, pb=True)
divider('f59e0b')
body('This version adds hosting architecture (Azure-first), full system requirements, and the future connector roadmap including SAP Business One integration.', color=C_GRAY)
tbl([
    ('New — §15', 'Hosting Architecture — Azure-first stack with GitHub CI/CD'),
    ('New — §16', 'System Requirements — development, production, browser, mobile'),
    ('New — §17', 'SAP Business One Integration — Service Layer API, daily/weekly dashboard data'),
    ('New — §18', 'Future Connector Roadmap — all planned integrations beyond MVP'),
    ('Updated — §1', 'Executive Summary updated to reflect Azure hosting and SAP roadmap'),
    ('No change', 'All sections from v2.0 (§2–§14) are carried forward unchanged'),
], col_w=(1.8,5.5), header='Changes from v2.0')


# ════════════════════════════════════════════════════════════
#  1. EXECUTIVE SUMMARY  (updated)
# ════════════════════════════════════════════════════════════
H('1. Executive Summary', 1, C_INDIGO, pb=True)
divider()
body(
    'TTD is a personal productivity system built around WhatsApp as its primary capture interface, '
    'with a three-panel Gemini web dashboard, Outlook calendar sync, and a full REST API backend. '
    'Infrastructure runs on Azure (the owner\'s existing cloud provider) with GitHub Actions for '
    'CI/CD. In future phases, TTD will connect to SAP Business One to pull live business data — '
    'sales orders, receivables, cash flow — into the same daily and weekly dashboard the owner '
    'already uses for personal tasks.'
)
tbl([
    ('Owner',             'Lior'),
    ('Primary capture',   'WhatsApp (text + voice)'),
    ('Web UI',            'Gemini — 3-panel dashboard (Kanban · Integrations · Timeline)'),
    ('Mailboxes',         'lior@lbatech.com  ·  liorba@mepsltn.com  (+extensible)'),
    ('AI stack',          'OpenAI GPT-4o  ·  Whisper STT'),
    ('Cloud provider',    'Microsoft Azure  (owner already has Azure subscription)'),
    ('Source control',    'GitHub  (owner already has GitHub account)'),
    ('CI/CD',             'GitHub Actions → Azure deployment pipelines'),
    ('ERP',               'SAP Business One — future connector for business dashboards'),
    ('Doc version',       '3.0  —  June 2026'),
], col_w=(2.0,5.3), header='Project Snapshot  v3.0')


# ════════════════════════════════════════════════════════════
#  SECTIONS 2–14: CARRIED FROM V2
# ════════════════════════════════════════════════════════════
H('2–14.  Sections Carried From v2.0', 1, C_GRAY, pb=True)
divider('888888')
note_box(
    'Sections 2 through 14 are unchanged from v2.0. They cover: System Architecture (§2), '
    'Gemini UI Component Inventory (§3), Data Model (§4), REST API (§5), Task Card JSON Contract (§6), '
    'Filter & View System (§7), External Integrations (§8), AI Agent Design (§9), Daily Report (§10), '
    'Security (§11), Build Phases (§12), Open Questions (§13), Glossary (§14). '
    'Refer to TTD_Product_Architecture_v2.0.docx for full detail on those sections.',
    fill='f1f5f9', border='c7d2fe'
)


# ════════════════════════════════════════════════════════════
#  15. HOSTING ARCHITECTURE
# ════════════════════════════════════════════════════════════
H('15. Hosting Architecture', 1, C_AZURE, pb=True)
divider('0078d4')
body(
    'The owner already holds an active Azure subscription and GitHub account. '
    'All hosting decisions below are Azure-first. Where a managed Supabase service is more '
    'practical for speed (Auth, Vault, Edge Functions), it is retained alongside Azure. '
    'Everything is containerised (Docker) so workloads can be migrated between providers without code changes.'
)

H('15.1  Service-to-Host Mapping', 2, C_AZURE)
tbl([
    ('FastAPI Backend',      'Azure Container Apps  —  auto-scales to zero, consumption billing, custom domain, TLS managed.  Alternatively: Azure App Service (B2/B3 plan) if always-on is preferred.'),
    ('React Web UI (Gemini)','Azure Static Web Apps  —  free tier available, GitHub Actions auto-deploy on push to main, global CDN built in, custom domain + TLS included.'),
    ('PostgreSQL Database',  'Azure Database for PostgreSQL — Flexible Server  (General Purpose, 2 vCore min).  OR  keep Supabase (simpler managed experience, built-in Auth + Vault).  Recommended: Supabase for MVP speed, migrate to Azure DB at scale.'),
    ('Auth & JWT',           'Supabase Auth  (email + OTP)  —  JWT tokens verified by FastAPI middleware.  Azure AD B2C is an alternative if SSO with company accounts is needed later.'),
    ('Secret Storage',       'Azure Key Vault  —  all API keys (OpenAI, Meta/WhatsApp, Twilio), OAuth tokens, DB connection strings.  FastAPI reads secrets at startup via Azure SDK.'),
    ('Cron / Scheduled Jobs','Azure Container Apps Jobs  —  runs daily report, token refresh, SAP sync on cron schedule.  OR  Supabase Edge Functions (simpler for small jobs).'),
    ('Object Storage',       'Azure Blob Storage  —  contact avatar uploads, voice note audio files (temporary), CSV import files.'),
    ('CDN',                  'Azure CDN (Front Door)  —  static assets for web UI, avatar images served from Blob Storage.  Already included in Static Web Apps.'),
    ('Monitoring & Logs',    'Azure Monitor + Application Insights  —  FastAPI request traces, error alerts, performance metrics.'),
    ('Source Control',       'GitHub  —  mono-repo: /backend (FastAPI) + /frontend (React).'),
    ('CI/CD',                'GitHub Actions  —  2 pipelines: backend (build Docker → push ACR → deploy Container Apps) + frontend (build → deploy Static Web Apps).'),
], col_w=(2.2,5.1), header='Service Hosting Map  (Azure-First)')

H('15.2  Azure Resource Group Layout', 2, C_AZURE)
body('All TTD resources live in a single Azure Resource Group for cost visibility and easy teardown.')
tbl([
    ('Resource Group',   'rg-ttdai-prod  (and rg-ttdai-dev for development environment)'),
    ('Container Apps Env','cae-ttdai-prod  —  hosts the FastAPI container app'),
    ('Container Registry','acr-ttdai  —  stores Docker images built by GitHub Actions'),
    ('Static Web App',   'swa-ttdai-frontend  —  Gemini React UI'),
    ('PostgreSQL Server', 'pg-ttdai-prod  (if using Azure DB)  OR  keep Supabase project'),
    ('Key Vault',        'kv-ttdai-prod  —  all secrets, rotated here not in code'),
    ('Storage Account',  'sttdaiprod  —  Blob containers: avatars, audio-temp, csv-imports'),
    ('App Insights',     'ai-ttdai-prod  —  connected to Container Apps for traces'),
], col_w=(2.2,5.1), header='Azure Resource Inventory')

H('15.3  CI/CD Pipeline — GitHub Actions', 2, C_AZURE)
body('Two workflow files in .github/workflows/:')
tbl([
    ('backend.yml',    'Trigger: push to main /backend/**. Steps: (1) Run tests, (2) Build Docker image, (3) Push to Azure Container Registry, (4) Deploy to Container Apps (zero-downtime rolling update).'),
    ('frontend.yml',   'Trigger: push to main /frontend/**. Steps: (1) npm install + build, (2) Deploy to Azure Static Web Apps via azure/static-web-apps-deploy action. Auto-generates preview URLs for PRs.'),
    ('Environments',   'dev branch → rg-ttdai-dev (auto-deploy). main branch → rg-ttdai-prod (requires manual approval gate in GitHub).'),
    ('Secrets in GH',  'AZURE_CREDENTIALS, ACR_LOGIN_SERVER, SWA_DEPLOYMENT_TOKEN stored as GitHub Secrets. Never in code.'),
], col_w=(2.2,5.1), header='GitHub Actions Pipelines')

H('15.4  Environment Strategy', 2, C_AZURE)
tbl([
    ('Development',  'Local Docker Compose (FastAPI + PostgreSQL + mock WA webhook). .env.local file. No Azure costs.'),
    ('Staging (dev)','Azure rg-ttdai-dev — real Azure services, test WhatsApp number, Supabase dev project. Auto-deployed from dev branch.'),
    ('Production',   'Azure rg-ttdai-prod — live WhatsApp number, production DB, Key Vault, manual deploy gate. All monitoring active.'),
], col_w=(1.8,5.5), header='Environments')

H('15.5  Cost Estimate (Monthly — Production)', 2, C_AZURE)
note_box('Estimates based on Azure pricing as of June 2026. Actual costs depend on usage volume.', fill='fff7ed', border='f59e0b')
tbl([
    ('Azure Container Apps',     '~$15–$30/mo  (consumption plan, scales to zero overnight)'),
    ('Azure Static Web Apps',    '$0  (Free tier covers TTD needs — 100GB bandwidth/mo)'),
    ('Azure Database for PostgreSQL', '~$50–$70/mo  (Flexible Server, 2 vCore, 8GB RAM)  OR  $0–$25 on Supabase free/pro'),
    ('Azure Key Vault',          '~$1–$3/mo  (10,000 operations/mo)'),
    ('Azure Blob Storage',       '~$2–$5/mo  (50GB storage + egress)'),
    ('Azure Container Registry', '~$5/mo  (Basic tier)'),
    ('Azure Application Insights','~$5–$15/mo  (5GB data ingestion)'),
    ('OpenAI API (GPT-4o + Whisper)', '~$30–$80/mo  (depends on WA message volume + voice notes)'),
    ('WhatsApp Business Cloud API','$0  (first 1,000 conversations/mo free, ~$0.005–$0.08 after)'),
    ('Supabase (if retained)',   '$0–$25/mo  (Free or Pro plan for Auth + Vault)'),
    ('Total estimate',           '~$110–$230/mo  depending on message volume and DB choice'),
], col_w=(2.8,4.5), header='Monthly Cost Estimate')


# ════════════════════════════════════════════════════════════
#  16. SYSTEM REQUIREMENTS
# ════════════════════════════════════════════════════════════
H('16. System Requirements', 1, C_INDIGO, pb=True)
divider()

H('16.1  Development Machine Requirements', 2, C_CYAN)
tbl([
    ('OS',          'Windows 10/11 (64-bit), macOS 12+, or Ubuntu 22.04+'),
    ('RAM',         '16 GB minimum  (32 GB recommended for running Docker + IDE simultaneously)'),
    ('CPU',         '4-core minimum  (8-core recommended)'),
    ('Disk',        '20 GB free space for Docker images, node_modules, Python venvs'),
    ('Docker',      'Docker Desktop 4.x+  (required for local dev environment)'),
    ('Node.js',     '20 LTS or 22 LTS  (for React frontend)'),
    ('Python',      '3.11 or 3.12  (for FastAPI backend)'),
    ('Git',         '2.40+  (GitHub CLI recommended)'),
    ('VS Code',     'Recommended IDE — Python + ESLint + Tailwind IntelliSense extensions'),
    ('Azure CLI',   'az CLI 2.x  (for deploying to Azure from terminal)'),
    ('GitHub CLI',  'gh CLI  (optional, for PR management)'),
], col_w=(2.0,5.3), header='Development Requirements')

H('16.2  Backend Runtime Requirements (Production)', 2, C_CYAN)
tbl([
    ('Runtime',         'Python 3.11+  inside Docker (python:3.11-slim base image)'),
    ('Container',       'Docker — FastAPI served via Uvicorn (4 workers minimum in production)'),
    ('vCPU',            '2 vCPU minimum  (Azure Container Apps auto-scales)'),
    ('RAM',             '1 GB minimum per container instance'),
    ('Database',        'PostgreSQL 15+  (Azure Flexible Server or Supabase)'),
    ('DB Storage',      '10 GB minimum  (grows with task_messages + audit_log volume)'),
    ('DB Connections',  'PgBouncer connection pooling recommended for Container Apps (max_connections=100)'),
    ('Outbound network','FastAPI must reach: api.openai.com, graph.microsoft.com, graph.facebook.com, api.twilio.com'),
    ('Inbound',         'HTTPS on port 443. WhatsApp webhook must be reachable on a public HTTPS URL.'),
    ('TLS',             'TLS 1.3 required. Azure Container Apps provides managed TLS.'),
    ('Python packages', 'fastapi, uvicorn[standard], sqlalchemy, asyncpg, httpx, openai, msal, python-jose, pydantic-settings'),
], col_w=(2.2,5.1), header='Backend Runtime (Production)')

H('16.3  Frontend Runtime Requirements (Production)', 2, C_CYAN)
tbl([
    ('Build output',    'Static files (HTML/CSS/JS) — no server runtime needed'),
    ('Hosting',         'Azure Static Web Apps (CDN-distributed)'),
    ('Build tool',      'Vite 5+  (faster than CRA, native ESM)'),
    ('Framework',       'React 18+  with TypeScript 5+'),
    ('Styling',         'Tailwind CSS 3.4+'),
    ('Bundle size',     'Target < 500 KB gzipped initial load'),
    ('CDN',             'Azure Front Door (bundled with Static Web Apps)'),
], col_w=(2.2,5.1), header='Frontend Runtime (Production)')

H('16.4  Browser Requirements', 2, C_CYAN)
tbl([
    ('Chrome',          '120+  (primary target — tested first)'),
    ('Edge',            '120+  (Chromium-based — second priority, owner likely uses this)'),
    ('Firefox',         '121+  (supported)'),
    ('Safari',          '17+  (Mac/iPad — supported)'),
    ('Mobile Safari',   'iOS 16+  (iPhone — glassmorphism backdrop-filter required)'),
    ('Not supported',   'IE 11, legacy Edge (EdgeHTML), Chrome < 100'),
    ('Required APIs',   'CSS backdrop-filter (glassmorphism), CSS Grid, Fetch API, WebSocket (future WA live feed)'),
    ('Screen resolution','1280×800 minimum. Optimised for 1440×900 and 1920×1080. 3-panel layout collapses to tabs on < 1024px wide.'),
], col_w=(2.2,5.1), header='Browser Compatibility')

H('16.5  Mobile Requirements', 2, C_CYAN)
body('WhatsApp is the mobile interface — the web dashboard is not intended as a mobile-first app. However:')
tbl([
    ('iOS',           'iPhone 12+ running iOS 16+. WhatsApp installed. Safari 17+ for web access.'),
    ('Android',       'Android 10+ with Chrome 120+. WhatsApp installed.'),
    ('Responsive',    'Web dashboard adapts: 3-panel → 2-panel at 1024px → single-panel tabs at 768px.'),
    ('WhatsApp',      'Primary mobile interface. Bot responses must be readable on 375px screens.'),
    ('PWA',           'Phase 2 optional: add PWA manifest + service worker so dashboard can be added to home screen.'),
], col_w=(2.0,5.3), header='Mobile Requirements')

H('16.6  WhatsApp Business Account Requirements', 2, C_CYAN)
tbl([
    ('Account type',    'WhatsApp Business Account (WABA) registered with Meta Business Manager'),
    ('Phone number',    'A dedicated phone number linked to the WABA. Can be the owner\'s existing number if migrated.'),
    ('API access',      'Meta Cloud API access — apply at developers.facebook.com. Approval typically 1–3 days.'),
    ('Webhook URL',     'A public HTTPS URL (Azure Container Apps provides this). Must respond to GET verification within 10 seconds.'),
    ('Message limits',  'Tier 1: 1,000 conversations/day free. Tier 2 (after 50 unique conversations): 10,000/day. No limit concern for personal use.'),
    ('Fallback',        'Twilio WhatsApp account as backup. Requires separate WABA or use of Twilio Sandbox for dev.'),
], col_w=(2.2,5.1), header='WhatsApp API Requirements')


# ════════════════════════════════════════════════════════════
#  17. SAP BUSINESS ONE INTEGRATION
# ════════════════════════════════════════════════════════════
H('17. SAP Business One Integration', 1, C_SAP, pb=True)
divider('009a44')
body(
    'The owner runs their companies on SAP Business One. The TTD system will connect to SAP B1 '
    'via its Service Layer REST API to pull business data into the daily and weekly dashboard — '
    'receivables, sales, cash position, open orders — alongside personal tasks in the same view.'
)
note_box(
    'SAP B1 Service Layer is available on SAP B1 version 9.3+ (on-premise) and SAP B1 Cloud. '
    'It exposes a REST/OData API running on HTTPS port 50000 (on-premise) or via SAP Cloud hostname. '
    'No SAP add-on installation is required — the Service Layer is built into B1.',
    fill='e8f5e9', border='009a44'
)

H('17.1  Connection Architecture', 2, C_SAP)
tbl([
    ('Protocol',        'SAP B1 Service Layer REST API (OData v4). HTTPS. Port 50000 (on-premise) or SAP cloud URL.'),
    ('Auth',            'Session-based: POST /b1s/v1/Login with CompanyDB, UserName, Password → returns SessionId cookie. Session valid 30 min. TTD backend refreshes automatically.'),
    ('Connection',      'FastAPI backend connects to SAP B1 Service Layer. Never from the browser — all SAP calls are server-side only.'),
    ('Frequency',       'Daily sync at 06:00 (before daily WhatsApp report). Weekly summary sync every Monday 06:00.'),
    ('On-demand',       'Owner can trigger a manual sync via WhatsApp command: "sap sync" or button in web dashboard.'),
    ('Data direction',  'READ ONLY from SAP B1 in all phases. Write-back (e.g., creating SAP tasks from TTD) is a future option.'),
    ('Credential storage','SAP username + password stored in Azure Key Vault. Never in code or DB.'),
], col_w=(2.2,5.1), header='SAP B1 Connection')

H('17.2  Daily Dashboard Data — SAP B1 Pulls', 2, C_SAP)
body('Data pulled every morning at 06:00 and included in the daily WhatsApp report and web dashboard.')
tbl([
    ('Open Sales Orders',       'GET /b1s/v1/Orders?$filter=DocumentStatus eq \'O\'  — count + total value of open orders'),
    ('Overdue Receivables (AR)', 'GET /b1s/v1/BusinessPartners with AccountReceivable balance > 0 and overdue > 0 days'),
    ('Cash & Bank Balance',     'GET /b1s/v1/ChartOfAccounts filtered to cash/bank accounts + current balance'),
    ('Invoices Due Today',      'GET /b1s/v1/Invoices?$filter=DocDueDate eq [today] and DocumentStatus eq \'O\''),
    ('Overdue Invoices',        'GET /b1s/v1/Invoices?$filter=DocDueDate lt [today] and DocumentStatus eq \'O\'  — top 5 by value'),
    ('Purchase Orders Pending', 'GET /b1s/v1/PurchaseOrders?$filter=DocumentStatus eq \'O\'  — awaiting delivery'),
    ('Low Stock Alerts',        'GET /b1s/v1/Items?$filter=QuantityOnStock lt MinimumInventory  — items below reorder point'),
], col_w=(2.8,4.5), header='Daily Sync — SAP B1 Endpoints')

H('17.3  Weekly Summary Data — SAP B1 Pulls', 2, C_SAP)
body('Pulled every Monday at 06:00. Added to Monday\'s WhatsApp report as "Weekly Business Summary".')
tbl([
    ('Week Revenue',            'GET /b1s/v1/Invoices with DocDate in last 7 days + sum of DocTotal'),
    ('New Customers',           'GET /b1s/v1/BusinessPartners?$filter=CreateDate ge [last_monday]  — new BPs added this week'),
    ('Top 5 Open Orders (value)','GET /b1s/v1/Orders sorted by DocTotal desc, top 5'),
    ('Payments Received (week)','GET /b1s/v1/IncomingPayments with DocDate in last 7 days'),
    ('Aged Receivables Summary','GET /b1s/v1/InvoicePaymentBalances  — 0-30 / 31-60 / 61-90 / 90+ day buckets'),
    ('Gross Profit (week)',     'Calculated from Invoices.GrossProfit field for the week period'),
], col_w=(2.8,4.5), header='Weekly Sync — SAP B1 Endpoints')

H('17.4  SAP Dashboard Data Model', 2, C_SAP)
body('SAP data is stored in a separate table — never mixed with personal tasks. Used only for dashboard widgets.')
tbl([
    ('sap_snapshots.id',         'UUID PRIMARY KEY'),
    ('sap_snapshots.company_db', 'TEXT — SAP CompanyDB name (supports multiple companies)'),
    ('sap_snapshots.snapshot_type', "TEXT — 'daily' | 'weekly'"),
    ('sap_snapshots.data',       'JSONB — full snapshot payload (receivables, orders, cash, etc.)'),
    ('sap_snapshots.synced_at',  'TIMESTAMPTZ — when this snapshot was pulled from SAP'),
    ('sap_snapshots.period_start','DATE — start of period this snapshot covers'),
    ('sap_snapshots.period_end', 'DATE — end of period'),
], col_w=(2.2,5.1), header='Table: sap_snapshots  (NEW)')

H('17.5  SAP Business Summary — WhatsApp Report Addition', 2, C_SAP)
body('When SAP sync is enabled, the daily WhatsApp report gains an extra section:')
tbl([
    ('Section header',       '📊 Business Summary — [CompanyDB name]'),
    ('Open orders',          '📦 Open Orders: N  (total value: $XXX,XXX)'),
    ('Cash position',        '💰 Cash & Bank: $XXX,XXX'),
    ('Overdue receivables',  '⚠️ Overdue AR: $XXX,XXX  (N invoices)'),
    ('Invoices due today',   '📋 Due today: N invoices  ($XXX,XXX)'),
    ('Weekly revenue',       '(Mondays only)  📈 This week revenue: $XXX,XXX'),
    ('Multi-company',        'If owner has multiple SAP company databases, each gets its own section with company name header.'),
], col_w=(2.2,5.1), header='SAP Section in Daily Report')

H('17.6  SAP Dashboard Widgets — Web UI', 2, C_SAP)
body('The Gemini web dashboard right panel gains a "Business" tab alongside the existing Outlook tabs:')
tbl([
    ('Business tab',        'New tab in right panel: "Business" — shows SAP snapshot data alongside Outlook Inbox / Tasks tabs.'),
    ('KPI cards',           'Cash & Bank · Open Orders Value · Overdue AR · Invoices Due Today — 4 metric cards at top of Business tab.'),
    ('AR aging bars',       'Horizontal bar chart: 0-30 / 31-60 / 61-90 / 90+ day receivable buckets.'),
    ('Top open orders',     'Table: order number, customer name, value, due date — top 5 by value.'),
    ('Last sync time',      '"Last updated: today 06:03" shown at top of Business tab.'),
    ('Manual sync button',  '"Sync now" button → POST /api/v1/sap/sync → triggers immediate SAP pull.'),
    ('Multi-company select','If multiple SAP company DBs: dropdown to switch between companies.'),
], col_w=(2.2,5.1), header='SAP Web Dashboard Widgets')

H('17.7  SAP B1 API Endpoints (TTD Backend)', 2, C_SAP)
tbl([
    ('GET  /api/v1/sap/snapshot/daily',   'Returns latest daily SAP snapshot for the dashboard.'),
    ('GET  /api/v1/sap/snapshot/weekly',  'Returns latest weekly SAP snapshot.'),
    ('POST /api/v1/sap/sync',             'Trigger immediate SAP sync (manual refresh). Returns new snapshot.'),
    ('GET  /api/v1/sap/companies',        'List configured SAP company databases.'),
    ('POST /api/v1/sap/companies',        'Add a new SAP company database connection (credentials stored in Key Vault).'),
    ('GET  /api/v1/sap/health',           'Check Service Layer connectivity + session status.'),
], col_w=(2.8,4.5), header='SAP API Endpoints (TTD)')

H('17.8  SAP Build Phase', 2, C_SAP)
body('SAP B1 integration is Phase 5 — after all core TTD phases are complete and stable.')
tbl([
    ('Phase 5 deliverables', 'sap_snapshots table · SAP Service Layer client (httpx session management) · daily + weekly cron sync · GET /sap/* endpoints · Business tab in web UI · SAP section in daily WhatsApp report · Multi-company support · Azure Key Vault SAP credentials'),
    ('Prerequisites',        'Phase 1–4 complete · SAP B1 Service Layer accessible from Azure (VPN or public Service Layer) · SAP user account with read-only API access · Network connectivity confirmed between Azure Container Apps and SAP server'),
    ('SAP network note',     'If SAP B1 is on-premise (not SAP Cloud), Azure must reach the SAP server. Options: (1) Azure VPN Gateway to office network, (2) SAP B1 Service Layer exposed via reverse proxy with IP whitelist, (3) SAP B1 Cloud — no networking needed.'),
], col_w=(2.2,5.1), header='SAP Phase 5 Plan')


# ════════════════════════════════════════════════════════════
#  18. FUTURE CONNECTOR ROADMAP
# ════════════════════════════════════════════════════════════
H('18. Future Connector Roadmap', 1, C_INDIGO, pb=True)
divider()
body('Connectors listed here are planned beyond the 5 core build phases. Each will follow the same pattern: server-side API client → sync cron → snapshot table → dashboard widget → WhatsApp report section.')

H('18.1  Connector Priority Matrix', 2, C_CYAN)
tbl([
    ('P0 — Core (Phases 1–4)', 'WhatsApp Cloud API · OpenAI (GPT-4o + Whisper) · Microsoft Outlook (Graph API) · Supabase Auth'),
    ('P1 — Phase 5',           'SAP Business One Service Layer  (daily/weekly business dashboard)'),
    ('P2 — Near-term',         'Google Calendar · Zoom / Teams Meetings · DocuSign (contract tracking) · Bank feeds (open banking)'),
    ('P3 — Mid-term',          'Salesforce CRM · HubSpot · Xero / QuickBooks (accounting) · Slack (team notifications)'),
    ('P4 — Long-term',         'SAP SuccessFactors (HR) · PowerBI embedded · Custom company intranet APIs'),
], col_w=(2.0,5.3), header='Connector Priority')

H('18.2  SAP Business One Extended — Phase 6+', 2, C_SAP)
body('After the initial read-only Phase 5, these write-back capabilities will be added:')
for item in [
    'Create SAP Activities from TTD tasks (POST /b1s/v1/Activities)',
    'Create SAP Service Calls from WhatsApp messages flagged as support issues',
    'Update SAP Business Partner notes from task completion in TTD',
    'Two-way sync: SAP Activities created in B1 appear as TTD tasks automatically',
    'SAP Crystal Reports triggered from TTD dashboard (weekly P&L, AR aging report)',
]:
    bullet(item)
sp()

H('18.3  Google Calendar', 2, C_CYAN)
tbl([
    ('Purpose',     'Sync meetings and tasks with Google Calendar alongside Outlook. For contacts who use Google.'),
    ('Auth',        'OAuth 2.0 — Google Calendar API v3. Same pattern as Outlook OAuth.'),
    ('Data',        'Read events, create events, send invites. Same meeting scheduler flow as Outlook.'),
    ('Phase',       'Phase 6  —  after Outlook integration is stable.'),
], col_w=(1.8,5.5), header='Google Calendar Connector')

H('18.4  Bank Feeds (Open Banking)', 2, C_CYAN)
tbl([
    ('Purpose',     'Pull daily bank balance and recent transactions into the Business tab alongside SAP data.'),
    ('API',         'Open Banking APIs (PSD2 in EU/UK) or bank-specific APIs. Provider TBD by owner.'),
    ('Data',        'Current balance per account · transactions last 7 days · upcoming payments.'),
    ('Phase',       'Phase 6+  —  depends on bank and open banking availability in owner\'s region.'),
], col_w=(1.8,5.5), header='Bank Feed Connector')

H('18.5  Zoom / Microsoft Teams', 2, C_CYAN)
tbl([
    ('Purpose',     'Meeting links auto-included in Outlook invites. Join meeting from TTD dashboard.'),
    ('Auth',        'Zoom: JWT or OAuth. Teams: already covered by MS Graph (Teams meetings via Graph).'),
    ('Data',        'Create Zoom meeting → attach link to TTD meeting record → include in Outlook invite.'),
    ('Phase',       'Phase 6  —  quick win once Outlook integration is live.'),
], col_w=(1.8,5.5), header='Zoom / Teams Connector')

H('18.6  Connector Architecture Pattern', 2, C_CYAN)
body('Every future connector follows this standard pattern to keep the codebase consistent:')
for step in [
    'Create connector module in /backend/connectors/{name}/ with: client.py (API client), sync.py (cron job), models.py (DB schema)',
    'Add {name}_snapshots table to PostgreSQL following the same pattern as sap_snapshots',
    'Register cron job in the scheduler (daily / weekly / on-demand)',
    'Add GET /api/v1/{name}/snapshot endpoint for the web UI',
    'Add new tab or widget to the Gemini web dashboard right panel',
    'Add connector section to daily WhatsApp report (only if owner enables it)',
    'Store all credentials in Azure Key Vault under /connectors/{name}/ path',
]:
    numbered(step)
sp()


# ════════════════════════════════════════════════════════════
#  19. OPEN QUESTIONS  (updated)
# ════════════════════════════════════════════════════════════
H('19. Open Questions  (v3)', 1, C_INDIGO, pb=True)
divider()
tbl([
    ('[P1 — Owner]  WA number',         'Use existing personal number (via Twilio Business API) or apply for a new dedicated Meta Cloud API number?'),
    ('[P1 — Owner]  Report time',        'Confirm daily report send time. Default: 07:00 Asia/Jerusalem.'),
    ('[P1 — Owner]  Bot language',       'Hebrew, English, or auto-mirror the sender language?'),
    ('[P2 — Owner]  Azure subscription','Confirm Azure subscription ID + region (Israel North / West Europe?) for resource deployment.'),
    ('[P2 — Owner]  Domain name',        'Confirm domain for web dashboard (CORS, TLS, Static Web Apps custom domain).'),
    ('[P2 — Gemini] API base URL',       'Confirm API base URL so Gemini UI environment variables can be set.'),
    ('[P3 — Owner]  Meeting duration',   'Default meeting duration when not stated: 30 min or 60 min?'),
    ('[P5 — Owner]  SAP network access', 'Is SAP B1 on-premise or SAP Cloud? If on-premise: does the office have a VPN or can Service Layer be exposed via reverse proxy with IP whitelist to Azure?'),
    ('[P5 — Owner]  SAP version',        'Confirm SAP B1 version (9.3+ required for Service Layer). Is Service Layer already enabled?'),
    ('[P5 — Owner]  SAP companies',      'How many company databases does TTD need to read from? Provide CompanyDB names.'),
    ('[P5 — Owner]  SAP read account',   'Create a dedicated read-only SAP user for TTD API access. Provide credentials (stored in Azure Key Vault — not shared in chat).'),
], col_w=(2.8,4.5), header='Open Questions  v3')


# ════════════════════════════════════════════════════════════
#  20. GLOSSARY  (updated)
# ════════════════════════════════════════════════════════════
H('20. Glossary', 1, C_INDIGO)
divider()
tbl([
    ('TTD AI',             'Task Tracker Dashboard — the product name'),
    ('SAP B1',          'SAP Business One — the ERP system the owner uses to run their companies'),
    ('Service Layer',   'SAP B1\'s built-in REST/OData API. Available on B1 version 9.3+. No add-on needed.'),
    ('sap_snapshots',   'PostgreSQL table storing periodic SAP data pulls as JSONB payloads'),
    ('Azure CAE',       'Azure Container Apps Environment — hosts the FastAPI Docker container'),
    ('Azure SWA',       'Azure Static Web Apps — hosts the Gemini React frontend with built-in CDN'),
    ('Azure KV',        'Azure Key Vault — encrypted secret store for all API keys and OAuth tokens'),
    ('ACR',             'Azure Container Registry — stores Docker images built by GitHub Actions'),
    ('rg-ttdai-prod',     'Azure Resource Group for production TTD resources'),
    ('Connector',       'A backend module that syncs data from an external system into TTD on a schedule'),
    ('L1–L5',           'Architecture layers: Ingestion, AI Processing, Domain, Storage, Delivery'),
    ('AI Agent',        'Background process reading inbound WA messages and proposing tasks'),
    ('due_status',      'Server-computed field: overdue | due_today | due_soon | upcoming | completed | no_date'),
    ('WA',              'WhatsApp'),
    ('P1 / P2 / P3',    'Task priority: 1=Business (red), 2=Personal (yellow), 3=General (green)'),
], col_w=(1.5,5.8), header='Glossary  v3')


# ── Save ─────────────────────────────────────────────────────
out = '/sessions/brave-kind-johnson/mnt/TTD/TTD_Product_Architecture_v3.0.docx'
doc.save(out)
print(f'Saved: {out}')
