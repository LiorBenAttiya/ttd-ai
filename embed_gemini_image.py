"""
embed_gemini_image.py
─────────────────────
Inserts the Gemini UI design image into TTD_Product_Architecture_v3.0.docx
as a full dedicated page immediately after the Change Log section.

HOW TO RUN:
  1. Copy your Gemini UI image to:
       C:\Users\lior\Claude\Projects\TTD\gemini_ui.png
     (any of these names work: gemini_ui.png / gemini_ui.jpg / gemini_ui.jpeg)
  2. Run:  python3 embed_gemini_image.py
  3. Output: TTD_Product_Architecture_v3.0_final.docx
"""

import os, sys
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Find image ────────────────────────────────────────────────
BASE   = os.path.dirname(os.path.abspath(__file__))
SRC    = os.path.join(BASE, 'TTD_Product_Architecture_v3.0.docx')
OUT    = os.path.join(BASE, 'TTD_Product_Architecture_v3.0_final.docx')

IMAGE_NAMES = ['gemini_ui.png', 'gemini_ui.jpg', 'gemini_ui.jpeg',
               'gemini_ui.webp', 'Gemini_UI.png', 'Gemini_UI.jpg']
img_path = None
for name in IMAGE_NAMES:
    candidate = os.path.join(BASE, name)
    if os.path.exists(candidate):
        img_path = candidate
        break

if not img_path:
    print("\n❌  Image not found. Please copy your Gemini UI design image to:")
    print(f"    {os.path.join(BASE, 'gemini_ui.png')}")
    print("\nThen re-run this script.")
    sys.exit(1)

print(f"✅  Found image: {img_path}")

# ── Load document ─────────────────────────────────────────────
doc = Document(SRC)

C_INDIGO = RGBColor(0x63,0x66,0xf1)
C_CYAN   = RGBColor(0x38,0xbd,0xf8)
C_WHITE  = RGBColor(0xff,0xff,0xff)
C_GRAY   = RGBColor(0x6b,0x72,0x80)
C_AMBER  = RGBColor(0xf5,0x9e,0x0b)

def shade_para(p, hex_fill):
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto')
    shd.set(qn('w:fill'), hex_fill)
    pPr.append(shd)

def page_break_para():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    br = OxmlElement('w:br'); br.set(qn('w:type'),'page')
    p.add_run()._r.append(br)
    return p

def center_para(text, size, color, bold=False, sb=0, sa=4):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(sb)
    p.paragraph_format.space_after  = Pt(sa)
    r = p.add_run(text); r.font.size = Pt(size)
    r.bold = bold; r.font.color.rgb = color
    return p

def divider_line(color_hex='6366f1'):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(8)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot  = OxmlElement('w:bottom')
    bot.set(qn('w:val'),'single'); bot.set(qn('w:sz'),'6')
    bot.set(qn('w:space'),'1');    bot.set(qn('w:color'), color_hex)
    pBdr.append(bot); pPr.append(pBdr)

# ── Build image page content ──────────────────────────────────
page_break_para()

# Section heading
h = doc.add_heading('UI Design Reference — Gemini Dashboard Blueprint', level=1)
h.paragraph_format.space_before = Pt(12)
h.paragraph_format.space_after  = Pt(4)
for run in h.runs: run.font.color.rgb = C_INDIGO

divider_line('6366f1')

# Caption above image
cap_top = doc.add_paragraph()
cap_top.alignment = WD_ALIGN_PARAGRAPH.CENTER
cap_top.paragraph_format.space_before = Pt(4)
cap_top.paragraph_format.space_after  = Pt(10)
r = cap_top.add_run('TTD AI — Three-Panel Dashboard  ·  "Ethereal Core" Dark Theme  ·  Designed by Gemini')
r.font.size = Pt(10); r.italic = True; r.font.color.rgb = C_GRAY

# ── Insert image centred, max width 6.3 inches ───────────────
try:
    from PIL import Image as PILImage
    with PILImage.open(img_path) as im:
        w_px, h_px = im.size
    aspect = h_px / w_px
    img_w  = Inches(6.3)
    img_h  = Inches(6.3 * aspect)
    # Cap height at 7 inches (leaves room for captions)
    if img_h > Inches(7):
        img_h = Inches(7)
        img_w = Inches(7 / aspect)
except ImportError:
    img_w = Inches(6.3)
    img_h = None   # let docx auto-size height

img_para = doc.add_paragraph()
img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
img_para.paragraph_format.space_before = Pt(0)
img_para.paragraph_format.space_after  = Pt(10)
run = img_para.add_run()
if img_h:
    run.add_picture(img_path, width=img_w, height=img_h)
else:
    run.add_picture(img_path, width=img_w)

# Caption below image
cap_bot = doc.add_paragraph()
cap_bot.alignment = WD_ALIGN_PARAGRAPH.CENTER
cap_bot.paragraph_format.space_before = Pt(4)
cap_bot.paragraph_format.space_after  = Pt(16)
r2 = cap_bot.add_run('Figure 1 — Gemini UI Design Blueprint (component layout + code architecture guide)')
r2.font.size = Pt(9); r2.italic = True; r2.font.color.rgb = C_GRAY

# Component legend table
from docx.enum.table import WD_TABLE_ALIGNMENT
legend_data = [
    ('Left Panel',   'Simple Kanban & Task Lists — To-Do · In-Progress · Done columns with progress bars, priority dots, source badges, and contact avatars'),
    ('Centre Panel', 'Integration Hub — Top: Outlook Inbox/Tasks (tabbed, with avatars + timestamps). Bottom: WhatsApp live message feed with green context badges'),
    ('Right Panel',  'Project Timeline & Analytics — Horizontal Gantt board + lifecycle blocks (To-Do / In-Progress / GA / Done) + Activity Stream'),
    ('Toolbar',      'Personal | Business toggle · List / Timeline / Kanban view switch · Filter dropdown · Due-date picker · Progress bar · Avatar row'),
    ('Theme',        '"Ethereal Core" — Deep charcoal/navy gradient (#0B0F19→#111827), glassmorphism surface cards, sage/dusty-rose/slate/ochre accent palette, sapphire geometric brand elements'),
    ('Components',   'card_task.tsx · column_kanban.tsx · integration_feed.tsx — React + Tailwind CSS, TypeScript, Vite 5'),
]

t = doc.add_table(rows=0, cols=2)
t.style = 'Table Grid'
t.alignment = WD_TABLE_ALIGNMENT.LEFT

def shade_cell(cell, hex_fill):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto'); shd.set(qn('w:fill'),hex_fill)
    tcPr.append(shd)

def cell_border(cell, color='c7d2fe'):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    b = OxmlElement('w:tcBorders')
    for side in ['top','left','bottom','right']:
        tg = OxmlElement(f'w:{side}')
        tg.set(qn('w:val'),'single'); tg.set(qn('w:sz'),'2')
        tg.set(qn('w:space'),'0');    tg.set(qn('w:color'),color)
        b.append(tg)
    tcPr.append(b)

# Header row
hr = t.add_row(); hr.cells[0].merge(hr.cells[1])
shade_cell(hr.cells[0], '1e1b4b'); cell_border(hr.cells[0])
hp = hr.cells[0].paragraphs[0]
run_h = hp.add_run('UI Component Breakdown')
run_h.bold = True; run_h.font.size = Pt(10); run_h.font.color.rgb = C_WHITE
hp.paragraph_format.space_before = Pt(4); hp.paragraph_format.space_after = Pt(4)
hp.paragraph_format.left_indent = Inches(0.1)

for left, right in legend_data:
    row = t.add_row()
    lc, rc = row.cells[0], row.cells[1]
    lc.width = Inches(1.6); rc.width = Inches(5.7)
    shade_cell(lc, 'f1f5f9'); shade_cell(rc, 'ffffff')
    cell_border(lc); cell_border(rc)
    lp = lc.paragraphs[0]
    lr = lp.add_run(left); lr.bold=True; lr.font.size=Pt(9); lr.font.color.rgb=C_INDIGO
    lp.paragraph_format.space_before=Pt(4); lp.paragraph_format.space_after=Pt(4)
    lp.paragraph_format.left_indent=Inches(0.1)
    rp = rc.paragraphs[0]
    rr = rp.add_run(right); rr.font.size=Pt(10)
    rp.paragraph_format.space_before=Pt(4); rp.paragraph_format.space_after=Pt(4)
    rp.paragraph_format.left_indent=Inches(0.1)

doc.add_paragraph().paragraph_format.space_after = Pt(8)

# ── Save ──────────────────────────────────────────────────────
doc.save(OUT)
print(f"\n✅  Saved: {OUT}")
print("   The Gemini UI design image has been embedded on its own page.")
