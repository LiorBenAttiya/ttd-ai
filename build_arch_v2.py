from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── Page setup ───────────────────────────────────────────────
s = doc.sections[0]
s.page_width = Inches(8.5); s.page_height = Inches(11)
s.left_margin = s.right_margin = Inches(1)
s.top_margin  = s.bottom_margin = Inches(1)

# ── Colours ──────────────────────────────────────────────────
C_INDIGO = RGBColor(0x63,0x66,0xf1)
C_CYAN   = RGBColor(0x38,0xbd,0xf8)
C_GREEN  = RGBColor(0x34,0xd3,0x99)
C_AMBER  = RGBColor(0xf5,0x9e,0x0b)
C_PINK   = RGBColor(0xc0,0x26,0xd3)
C_RED    = RGBColor(0xef,0x44,0x44)
C_GRAY   = RGBColor(0x6b,0x72,0x80)
C_WHITE  = RGBColor(0xff,0xff,0xff)
C_DARK   = RGBColor(0x1a,0x1a,0x2e)

# ── Helpers ──────────────────────────────────────────────────
def shade_para(p, hex_fill):
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto'); shd.set(qn('w:fill'),hex_fill)
    pPr.append(shd)

def shade_cell(cell, hex_fill):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto'); shd.set(qn('w:fill'),hex_fill)
    tcPr.append(shd)

def cell_border(cell, color='c7d2fe'):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    b = OxmlElement('w:tcBorders')
    for side in ['top','left','bottom','right']:
        t = OxmlElement(f'w:{side}')
        t.set(qn('w:val'),'single'); t.set(qn('w:sz'),'2')
        t.set(qn('w:space'),'0');    t.set(qn('w:color'),color)
        b.append(t)
    tcPr.append(b)

def divider(color='6366f1'):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'),'single'); bot.set(qn('w:sz'),'6')
    bot.set(qn('w:space'),'1');    bot.set(qn('w:color'),color)
    pBdr.append(bot); pPr.append(pBdr)

def H(text, level=1, color=None, pb=False):
    if pb:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0)
        br = OxmlElement('w:br'); br.set(qn('w:type'),'page')
        p.add_run()._r.append(br)
    h = doc.add_heading(text, level=level)
    h.paragraph_format.space_before = Pt(12 if level==1 else 8)
    h.paragraph_format.space_after  = Pt(4)
    if color:
        for r in h.runs: r.font.color.rgb = color
    return h

def body(text, bold=False, color=None, size=10, after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(after)
    r = p.add_run(text); r.font.size = Pt(size); r.bold = bold
    if color: r.font.color.rgb = color
    return p

def bullet(text, color=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(1); p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Inches(0.25)
    r = p.add_run(text); r.font.size = Pt(10)
    if color: r.font.color.rgb = color

def numbered(text, color=None):
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.space_before = Pt(1); p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Inches(0.25)
    r = p.add_run(text); r.font.size = Pt(10)
    if color: r.font.color.rgb = color

def tbl(rows, col_w=(2.0,5.3), header=None, hfill='1e1b4b'):
    t = doc.add_table(rows=0, cols=2); t.style='Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    bc = 'c7d2fe'
    if header:
        hr = t.add_row(); hr.cells[0].merge(hr.cells[1])
        shade_cell(hr.cells[0], hfill); cell_border(hr.cells[0], bc)
        hp = hr.cells[0].paragraphs[0]
        r = hp.add_run(header); r.bold=True; r.font.size=Pt(10); r.font.color.rgb=C_WHITE
        hp.paragraph_format.space_before=Pt(4); hp.paragraph_format.space_after=Pt(4)
        hp.paragraph_format.left_indent=Inches(0.1)
    for left, right in rows:
        row = t.add_row(); lc,rc = row.cells[0],row.cells[1]
        lc.width=Inches(col_w[0]); rc.width=Inches(col_w[1])
        shade_cell(lc,'f1f5f9'); shade_cell(rc,'ffffff')
        cell_border(lc,bc); cell_border(rc,bc)
        lp=lc.paragraphs[0]; lr=lp.add_run(left)
        lr.bold=True; lr.font.size=Pt(9); lr.font.color.rgb=C_INDIGO
        lp.paragraph_format.space_before=Pt(4); lp.paragraph_format.space_after=Pt(4)
        lp.paragraph_format.left_indent=Inches(0.1)
        rp=rc.paragraphs[0]; rr=rp.add_run(right); rr.font.size=Pt(10)
        rp.paragraph_format.space_before=Pt(4); rp.paragraph_format.space_after=Pt(4)
        rp.paragraph_format.left_indent=Inches(0.1)
    doc.add_paragraph().paragraph_format.space_after=Pt(4)

def sp(): doc.add_paragraph().paragraph_format.space_after = Pt(4)

# ════════════════════════════════════════════════════════════
#  TITLE PAGE
# ════════════════════════════════════════════════════════════
for text, size, color, sb, sa in [
    ('PRODUCT ARCHITECTURE DOCUMENT', 22, C_WHITE,  48, 4),
    ('Personal Task Tracker — TTD',   16, C_CYAN,    0, 4),
    ('WhatsApp-First  ·  AI Agent  ·  Voice  ·  Meetings  ·  Multi-Mailbox  ·  Web', 10, C_GRAY, 0, 2),
    ('Version 2.0  —  June 2026  —  Updated with Gemini UI Design', 9, C_AMBER, 0, 2),
    ('Owner: Lior   |   UI: Gemini   |   Backend: FastAPI + Supabase', 9, C_GRAY, 0, 60),
]:
    p = doc.add_paragraph(); shade_para(p,'1a1a2e')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before=Pt(sb); p.paragraph_format.space_after=Pt(sa)
    r = p.add_run(text); r.bold=(size>=16); r.font.size=Pt(size); r.font.color.rgb=color


# ════════════════════════════════════════════════════════════
#  CHANGE LOG v1 → v2
# ════════════════════════════════════════════════════════════
H('Change Log  —  v1.0 → v2.0', 1, C_AMBER, pb=True)
divider('f59e0b')
body('This version incorporates the complete Gemini UI design delivered for TTD. Every new API endpoint, data field, and architectural decision below is derived directly from the UI component inventory.', color=C_GRAY)
tbl([
    ('New section',     '§3   — Gemini UI Component Inventory (full panel-by-panel breakdown)'),
    ('New section',     '§5   — Dashboard API Endpoints (new endpoints required by the UI)'),
    ('Updated',         '§6   — Data Model: 8 new fields across tasks, contacts, mailboxes tables'),
    ('Updated',         '§7   — REST API: 9 new endpoints added (stats, timeline, inbox, avatars)'),
    ('New section',     '§11  — Task Card Contract (exact JSON shape the UI expects per card)'),
    ('Updated',         '§12  — Filter & View System (Personal/Business toggle, List/Timeline toggle)'),
    ('No change',       '§1,2,4,8,9,10,13,14,15 — unchanged from v1.0'),
], col_w=(2.0,5.3), header='Changes from v1.0')


# ════════════════════════════════════════════════════════════
#  1. EXECUTIVE SUMMARY
# ════════════════════════════════════════════════════════════
H('1. Executive Summary', 1, C_INDIGO, pb=True)
divider()
body(
    'TTD (Task Tracker Dashboard) is a personal productivity system with WhatsApp as its primary '
    'capture interface. The Gemini-designed web dashboard is a three-panel layout: a Kanban board '
    '(left), an integrations hub showing WhatsApp threads and activity (centre), and a Project '
    'Timeline + Outlook Inbox panel (right). The backend is API-first — FastAPI + PostgreSQL '
    '(Supabase) — and the Gemini UI consumes all data through clean REST endpoints.'
)
tbl([
    ('Owner',            'Lior'),
    ('Primary capture',  'WhatsApp (text + voice)'),
    ('Web UI',           'Gemini — 3-panel dashboard (Kanban · Integrations · Timeline)'),
    ('Mailboxes',        'lior@lbatech.com  ·  liorba@mepsltn.com  (+extensible)'),
    ('AI stack',         'OpenAI GPT-4o · Whisper STT'),
    ('Backend',          'Python FastAPI + PostgreSQL (Supabase)'),
    ('Messaging',        'WhatsApp Business Cloud API / Twilio fallback'),
    ('Doc version',      '2.0  —  June 2026  (UI-aligned)'),
], col_w=(2.0,5.3), header='Project Snapshot')


# ════════════════════════════════════════════════════════════
#  2. SYSTEM ARCHITECTURE
# ════════════════════════════════════════════════════════════
H('2. System Architecture — 5-Layer Model', 1, C_INDIGO, pb=True)
divider()
tbl([
    ('L1  Ingestion',    'WhatsApp webhooks (text + audio), Outlook Graph webhooks, REST calls from Gemini UI.'),
    ('L2  AI Processing','Whisper transcription, GPT-4o NLP, intent classification, meeting parsing, AI agent.'),
    ('L3  Domain',       'Task CRUD, approval flow, meeting scheduling, contact matching, report generation, mailbox routing.'),
    ('L4  Storage',      'PostgreSQL/Supabase: tasks, mailboxes, contacts, companies, meetings, approval_queue, dashboard_stats, audit_log.'),
    ('L5  Delivery',     'WhatsApp outbound, Outlook calendar invites, REST API for Gemini UI (all 3 panels + toolbar).'),
], col_w=(1.8,5.5), header='Architecture Layers')

H('2.1  Component Tree (from Gemini architecture guide)', 2, C_CYAN)
body('The Gemini UI declares the following component hierarchy. Every node maps to one or more backend API calls.')
tbl([
    ('DashboardView',        'Root component. Loads dashboard stats on mount. Hosts toolbar and all 3 panels.'),
    ('  LeftPanel [Kanban]', 'Fetches tasks grouped by status (pending / in_progress / completed). Renders draggable cards.'),
    ('  CenterPanel [Integrations]', 'Fetches linked WA message threads and activity feed for the selected task.'),
    ('  RightPanel [Timeline]',      'Two sub-views: (a) Project Timeline / Gantt, (b) Outlook Inbox & Tasks tabs.'),
    ('Integration Services', 'Singleton layer: WhatsApp connector, Outlook connector, AI agent status.'),
], col_w=(2.6,4.7), header='Gemini Component Tree')


# ════════════════════════════════════════════════════════════
#  3. GEMINI UI — COMPONENT INVENTORY
# ════════════════════════════════════════════════════════════
H('3. Gemini UI — Component Inventory', 1, C_INDIGO, pb=True)
divider()
body('Full panel-by-panel breakdown of every UI element observed in the Gemini design. Each row lists the component, what data it needs, and the API call that serves it.')

H('3.1  Toolbar / Filter Bar', 2, C_CYAN)
tbl([
    ('View toggle',         '"List View" | "Timeline View" — switches LeftPanel between Kanban columns and flat list.  →  ?view=kanban|list|timeline'),
    ('Category toggle',     '"Personal" | "Business" — maps to priority filter.  →  ?category=personal|business|all'),
    ('Filter dropdown',     'Multi-select: status, priority, source, mailbox, due date range.  →  query params on GET /tasks'),
    ('Due-date picker',     '"Due on..." filter. Sends due_before / due_after params.'),
    ('Progress bar',        'Overall completion %. Pulled from GET /api/v1/dashboard/stats → { total, completed, pct }.'),
    ('Avatar row',          'Contact avatars of people involved in recent tasks. Pulled from GET /api/v1/contacts?recent=true.'),
], col_w=(2.2,5.1), header='Toolbar Components')

H('3.2  Left Panel — Kanban Board', 2, C_CYAN)
tbl([
    ('Column headers',      'To-Do  ·  In-Progress  ·  Done. Each shows task count badge.'),
    ('Task card — title',   'tasks.description (truncated to ~60 chars)'),
    ('Task card — body',    'tasks.description (full snippet below title)'),
    ('Task card — priority','Colour dot: 🔴 red border = P1, 🟡 yellow = P2, 🟢 green = P3'),
    ('Task card — due badge','"Due Today" (orange) / "Got Delay" / "Overdue" — derived from due_date vs today'),
    ('Task card — source icon','WhatsApp logo or Outlook logo — from tasks.source field'),
    ('Task card — avatar',  'contacts.avatar_url of linked contact. Fallback: initials.'),
    ('Task card — IN badge','tasks.mailbox_id shortname: IN = inbox source flag'),
    ('Drag-to-reorder',     'PATCH /api/v1/tasks/{id}  with {status: "in_progress"} on column drop'),
], col_w=(2.2,5.1), header='Kanban Components')

H('3.3  Centre Panel — Integrations Hub', 2, C_CYAN)
tbl([
    ('WA thread preview',   'GET /api/v1/tasks/{id}/messages — returns last 10 WA message snippets linked to this task'),
    ('Activity feed',       'GET /api/v1/tasks/{id}/activity — audit log entries for this task (created, edited, completed)'),
    ('Contact card',        'GET /api/v1/contacts/{id} — name, role, company, avatar, phone, whatsapp, email'),
    ('Source label',        'tasks.source rendered as human label: "From WhatsApp" | "From Outlook" | "Voice Note" | "AI Proposed"'),
], col_w=(2.2,5.1), header='Centre Panel Components')

H('3.4  Right Panel — Two-Tab View', 2, C_CYAN)
tbl([
    ('Tab A: Outlook Inbox', 'GET /api/v1/mailboxes/{id}/inbox — returns unread emails with subject, sender, timestamp. Badge shows unread count.'),
    ('Tab B: Outlook Tasks', 'GET /api/v1/tasks?source=outlook_flag — tasks flagged from Outlook. Shows overdue count badge.'),
    ('Inbox row',           'subject, sender_name, sender_email, received_at, is_read, flag_status'),
    ('Task row',            'description, due_date, due_status (due_today / overdue / upcoming), source_icon, mailbox shortname'),
    ('Tab badges',          '"Inbox (5 unread)"  →  count from mailbox unread_count field. "Tasks (9 overdue)"  →  count from overdue tasks query.'),
], col_w=(2.2,5.1), header='Right Panel — Outlook Tabs')

H('3.5  Right Panel — Project Timeline (Gantt)', 2, C_CYAN)
tbl([
    ('Timeline rows',       'One row per task or project group. GET /api/v1/tasks?view=timeline returns tasks with start_date + due_date.'),
    ('Date-range blocks',   'Coloured bar from start_date to due_date. Colour = priority (red/yellow/green).'),
    ('Column headers',      'Date range (week/month). Configurable via ?timeline_from + ?timeline_to params.'),
    ('Item labels',         'tasks.description (short). Clicking opens centre panel with task detail.'),
    ('Scroll / zoom',       'Frontend-only — no additional API calls. Data already loaded from timeline endpoint.'),
], col_w=(2.2,5.1), header='Timeline / Gantt Components')

H('3.6  Status & Due Badges (derived fields)', 2, C_CYAN)
body('These are computed by the API and returned as a due_status field on every task, so the UI never has to calculate dates:')
tbl([
    ('overdue',     'due_date < today AND status != completed'),
    ('due_today',   'due_date = today AND status != completed'),
    ('due_soon',    'due_date within next 3 days AND status != completed'),
    ('upcoming',    'due_date > today+3 AND status != completed'),
    ('completed',   'status = completed'),
    ('no_date',     'due_date IS NULL'),
], col_w=(1.5,5.8), header='due_status Values (computed by API)')


# ════════════════════════════════════════════════════════════
#  4. DATA MODEL  (updated)
# ════════════════════════════════════════════════════════════
H('4. Data Model  (v2 — UI-Aligned)', 1, C_INDIGO, pb=True)
divider()
body('Fields marked  ★ NEW  were added in v2 based on Gemini UI requirements.')

H('4.1  tasks  (updated)', 2, C_CYAN)
tbl([
    ('id',            'UUID  PRIMARY KEY  DEFAULT gen_random_uuid()'),
    ('description',   'TEXT  NOT NULL'),
    ('due_date',      'DATE'),
    ('start_date',    'DATE  NULLABLE  ★ NEW — for Gantt / Timeline view'),
    ('priority',      'SMALLINT  NOT NULL  CHECK (priority IN (1,2,3))  — 1=Business  2=Personal  3=General'),
    ('status',        "TEXT  NOT NULL  DEFAULT 'pending'  CHECK (status IN ('pending','in_progress','completed','archived'))"),
    ('due_status',    "TEXT  GENERATED  — 'overdue'|'due_today'|'due_soon'|'upcoming'|'completed'|'no_date'  ★ NEW"),
    ('source',        "TEXT  — 'whatsapp_text'|'whatsapp_voice'|'web'|'outlook_flag'|'ai_agent'"),
    ('source_label',  "TEXT  GENERATED  — human-readable: 'From WhatsApp'|'Voice Note'|'From Outlook'|'AI Proposed'  ★ NEW"),
    ('mailbox_id',    'UUID  FK → mailboxes.id  NULLABLE'),
    ('contact_id',    'UUID  FK → contacts.id   NULLABLE'),
    ('meeting_id',    'UUID  FK → meetings.id   NULLABLE'),
    ('sort_order',    'INTEGER  DEFAULT 0  ★ NEW — Kanban drag-to-reorder position within column'),
    ('created_at',    'TIMESTAMPTZ  DEFAULT now()'),
    ('completed_at',  'TIMESTAMPTZ  NULLABLE'),
    ('deleted_at',    'TIMESTAMPTZ  NULLABLE  — soft delete'),
], col_w=(2.0,5.3), header='Table: tasks')

H('4.2  contacts  (updated)', 2, C_CYAN)
tbl([
    ('id',          'UUID  PRIMARY KEY'),
    ('company_id',  'UUID  FK → companies.id'),
    ('full_name',   'TEXT  NOT NULL'),
    ('email',       'TEXT'),
    ('phone',       'TEXT'),
    ('whatsapp',    'TEXT'),
    ('role',        'TEXT'),
    ('avatar_url',  'TEXT  NULLABLE  ★ NEW — photo URL for Kanban card avatar'),
    ('initials',    'TEXT  GENERATED  ★ NEW — fallback when no avatar (e.g., "LB")'),
    ('notes',       'TEXT'),
    ('created_at',  'TIMESTAMPTZ  DEFAULT now()'),
], col_w=(2.0,5.3), header='Table: contacts')

H('4.3  mailboxes  (updated)', 2, C_CYAN)
tbl([
    ('id',                  'UUID  PRIMARY KEY'),
    ('email',               'TEXT  UNIQUE  NOT NULL'),
    ('display_name',        'TEXT  NOT NULL'),
    ('short_name',          "TEXT  NOT NULL  ★ NEW — e.g., 'lbatech' | 'mepsltn' — shown as IN badge on cards"),
    ('unread_count',        'INTEGER  DEFAULT 0  ★ NEW — cached unread email count for Inbox tab badge'),
    ('last_synced_at',      'TIMESTAMPTZ  NULLABLE  ★ NEW — last Graph API inbox sync'),
    ('oauth_access_token',  'TEXT  ENCRYPTED'),
    ('oauth_refresh_token', 'TEXT  ENCRYPTED'),
    ('token_expires_at',    'TIMESTAMPTZ'),
    ('active',              'BOOLEAN  DEFAULT TRUE'),
    ('created_at',          'TIMESTAMPTZ  DEFAULT now()'),
], col_w=(2.0,5.3), header='Table: mailboxes')

H('4.4  task_messages  (NEW table)', 2, C_CYAN)
body('★ NEW — stores WhatsApp message snippets linked to tasks for the Centre Panel thread view.')
tbl([
    ('id',         'UUID  PRIMARY KEY'),
    ('task_id',    'UUID  FK → tasks.id  NOT NULL'),
    ('wa_msg_id',  'TEXT  — original WhatsApp message ID'),
    ('direction',  "TEXT  CHECK (direction IN ('inbound','outbound'))"),
    ('body',       'TEXT  NOT NULL'),
    ('sent_at',    'TIMESTAMPTZ  NOT NULL'),
    ('created_at', 'TIMESTAMPTZ  DEFAULT now()'),
], col_w=(2.0,5.3), header='Table: task_messages  (NEW)')

H('4.5  task_activity  (NEW table)', 2, C_CYAN)
body('★ NEW — audit/activity log per task for the Centre Panel activity feed.')
tbl([
    ('id',          'UUID  PRIMARY KEY'),
    ('task_id',     'UUID  FK → tasks.id  NOT NULL'),
    ('action',      "TEXT  — 'created'|'updated'|'completed'|'approved'|'rejected'|'meeting_linked'"),
    ('detail',      'JSONB  — {field, old_value, new_value}  NULLABLE'),
    ('created_at',  'TIMESTAMPTZ  DEFAULT now()'),
], col_w=(2.0,5.3), header='Table: task_activity  (NEW)')

H('4.6  Unchanged tables (v1)', 2, C_CYAN)
body('The following tables are unchanged from v1.0: companies, meetings, approval_queue, scheduled_reports.')


# ════════════════════════════════════════════════════════════
#  5. REST API  (v2 — full endpoint list)
# ════════════════════════════════════════════════════════════
H('5. REST API  (v2 — Complete Endpoint List)', 1, C_INDIGO, pb=True)
divider()
body('Base path: /api/v1. Auth: Bearer JWT. Endpoints marked  ★ NEW  were added in v2 for Gemini UI.')

H('5.1  Dashboard (NEW)', 2, C_CYAN)
tbl([
    ('GET  /dashboard/stats',        '★ NEW  —  Progress bar data. Returns { total, pending, in_progress, completed, overdue, pct_complete }.'),
    ('GET  /dashboard/recent-contacts','★ NEW  —  Avatar row in toolbar. Returns last 5 contacts linked to recent tasks.'),
    ('GET  /dashboard/activity',     '★ NEW  —  Global activity feed for centre panel default state (no task selected).'),
], col_w=(2.8,4.5), header='Dashboard Endpoints  ★ NEW')

H('5.2  Tasks (updated)', 2, C_CYAN)
tbl([
    ('GET    /tasks',               'List tasks. Params: status, priority, source, mailbox_id, category (personal|business|all), view (kanban|list|timeline), due_status, due_before, due_after, timeline_from, timeline_to, limit, offset.'),
    ('POST   /tasks',               'Create task. Body: {description, due_date, start_date?, priority, mailbox_id?, contact_id?}'),
    ('GET    /tasks/{id}',          'Single task with contact, meeting, mailbox, due_status, source_label.'),
    ('PATCH  /tasks/{id}',          'Partial update. Includes sort_order for Kanban drag-drop.'),
    ('DELETE /tasks/{id}',          'Soft delete.'),
    ('POST   /tasks/{id}/complete', 'Set completed.'),
    ('GET    /tasks/{id}/messages', '★ NEW  —  Centre panel WA thread. Returns task_messages[] ordered by sent_at.'),
    ('GET    /tasks/{id}/activity', '★ NEW  —  Centre panel activity feed. Returns task_activity[] ordered by created_at desc.'),
    ('PATCH  /tasks/reorder',       '★ NEW  —  Kanban drag-drop. Body: [{id, status, sort_order}]. Batch update positions.'),
], col_w=(2.8,4.5), header='Task Endpoints')

H('5.3  Outlook Inbox (NEW)', 2, C_CYAN)
tbl([
    ('GET  /mailboxes/{id}/inbox',      '★ NEW  —  Right panel Inbox tab. Returns emails[{subject, sender_name, sender_email, received_at, is_read, flag_status}] + unread_count.'),
    ('POST /mailboxes/{id}/inbox/sync', '★ NEW  —  Force refresh inbox from MS Graph. Updates unread_count on mailboxes table.'),
    ('POST /mailboxes/{id}/inbox/{email_id}/flag', '★ NEW  —  Flag an email from inbox panel → creates task. Returns new task_id.'),
], col_w=(2.8,4.5), header='Outlook Inbox Endpoints  ★ NEW')

H('5.4  Timeline (NEW)', 2, C_CYAN)
tbl([
    ('GET  /tasks?view=timeline', '★ NEW (param)  —  Returns tasks with start_date + due_date + priority for Gantt rendering. Required params: timeline_from, timeline_to.'),
], col_w=(2.8,4.5), header='Timeline / Gantt Endpoint')

H('5.5  Existing endpoints (unchanged from v1)', 2, C_CYAN)
body('The following endpoint groups are unchanged: /approvals, /meetings, /companies, /contacts, /reports/*, /mailboxes (connect/disconnect), /webhook/*. See v1.0 for full specs.')


# ════════════════════════════════════════════════════════════
#  6. TASK CARD JSON CONTRACT
# ════════════════════════════════════════════════════════════
H('6. Task Card JSON Contract', 1, C_INDIGO, pb=True)
divider()
body('This is the exact JSON shape GET /api/v1/tasks must return per item. The Gemini UI reads every field listed here.')
body('Example response object:', bold=True)

code_lines = [
    '{',
    '  "id":           "3f2a...",',
    '  "description":  "Call investor re: Series A",',
    '  "due_date":     "2026-06-04",',
    '  "start_date":   "2026-06-01",',
    '  "priority":     1,',
    '  "priority_label": "Business",',
    '  "priority_color": "#ef4444",',
    '  "status":       "pending",',
    '  "due_status":   "due_today",',
    '  "due_badge":    "Due Today",',
    '  "due_badge_color": "#f59e0b",',
    '  "source":       "whatsapp_voice",',
    '  "source_label": "Voice Note",',
    '  "source_icon":  "whatsapp",',
    '  "sort_order":   3,',
    '  "mailbox": {',
    '    "id":         "...",',
    '    "short_name": "lbatech",',
    '    "email":      "lior@lbatech.com"',
    '  },',
    '  "contact": {',
    '    "id":         "...",',
    '    "full_name":  "David Cohen",',
    '    "avatar_url": "https://...",',
    '    "initials":   "DC"',
    '  },',
    '  "created_at":   "2026-06-03T14:22:00Z",',
    '  "completed_at": null',
    '}',
]
for line in code_lines:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.left_indent  = Inches(0.3)
    r = p.add_run(line); r.font.size = Pt(9); r.font.name = 'Courier New'
    r.font.color.rgb = RGBColor(0x22,0x22,0x44)
sp()

tbl([
    ('priority_color',   'Hex string consumed directly by UI for dot/border colour. Server computes from priority int.'),
    ('due_status',       'Enum consumed by UI for badge selection. Never calculated in frontend.'),
    ('due_badge',        'Human-readable badge text: "Due Today" | "Overdue" | "Due Soon" | "Upcoming".'),
    ('due_badge_color',  'Hex for badge background. Amber = due_today, Red = overdue, Blue = upcoming.'),
    ('source_icon',      '"whatsapp" | "outlook" | "web" | "voice" — UI maps to icon component.'),
    ('contact.initials', 'Two-letter fallback when avatar_url is null. Computed from full_name.'),
], col_w=(2.0,5.3), header='Computed Fields  —  Why They Exist')


# ════════════════════════════════════════════════════════════
#  7. FILTER & VIEW SYSTEM
# ════════════════════════════════════════════════════════════
H('7. Filter & View System', 1, C_INDIGO, pb=True)
divider()
body('The toolbar in the Gemini UI drives all filtering via query parameters on GET /api/v1/tasks.')

H('7.1  Category Toggle  (Personal | Business)', 2, C_CYAN)
tbl([
    ('?category=business',  'Returns tasks where priority = 1'),
    ('?category=personal',  'Returns tasks where priority = 2'),
    ('?category=all',       'No priority filter (default)'),
], col_w=(2.5,4.8), header='Category Filter')

H('7.2  View Toggle  (Kanban | List | Timeline)', 2, C_CYAN)
tbl([
    ('?view=kanban',    'Default. Returns tasks grouped by status in response envelope: { todo:[], in_progress:[], done:[] }'),
    ('?view=list',      'Returns flat array ordered by due_date asc, then priority asc.'),
    ('?view=timeline',  'Returns flat array ordered by start_date asc. Requires timeline_from + timeline_to params. Includes start_date + due_date for Gantt bars.'),
], col_w=(2.5,4.8), header='View Modes')

H('7.3  Status Filter Mapping', 2, C_CYAN)
tbl([
    ('Kanban "To-Do" column',         "?status=pending"),
    ('Kanban "In-Progress" column',   "?status=in_progress"),
    ('Kanban "Done" column',          "?status=completed"),
    ('Right panel "Tasks (overdue)"', "?due_status=overdue"),
    ('Report: tasks due today',       "?due_status=due_today"),
    ('Report: next 7 days',           "?due_before=[today+7]&status=pending,in_progress"),
], col_w=(2.8,4.5), header='Status-to-Query Mapping')

H('7.4  Kanban Response Envelope  (view=kanban)', 2, C_CYAN)
kanban_lines = [
    '{',
    '  "todo":        [ /* tasks where status=pending, ordered by sort_order */ ],',
    '  "in_progress": [ /* tasks where status=in_progress, ordered by sort_order */ ],',
    '  "done":        [ /* tasks where status=completed, ordered by completed_at desc */ ],',
    '  "counts": {',
    '    "todo": 4,',
    '    "in_progress": 2,',
    '    "done": 11',
    '  }',
    '}',
]
for line in kanban_lines:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.left_indent = Inches(0.3)
    r = p.add_run(line); r.font.size = Pt(9); r.font.name = 'Courier New'
    r.font.color.rgb = RGBColor(0x22,0x22,0x44)
sp()


# ════════════════════════════════════════════════════════════
#  8. INTEGRATIONS  (unchanged summary)
# ════════════════════════════════════════════════════════════
H('8. External Integrations  (Summary)', 1, C_INDIGO, pb=True)
divider()
body('Full detail in v1.0. Summary retained here for completeness.')
tbl([
    ('WhatsApp Cloud API', 'Inbound webhooks + outbound messages. Audio → Whisper pipeline. HMAC-verified.'),
    ('OpenAI GPT-4o',      'NLP task parsing, AI agent intent classification, meeting extraction, confidence scoring.'),
    ('OpenAI Whisper',     'Speech-to-text for voice notes. Auto-detects Hebrew and English.'),
    ('MS Graph API',       'Outlook calendar event CRUD + mail inbox read + flag change subscriptions. Multi-mailbox OAuth.'),
    ('Twilio',             'Fallback WhatsApp delivery on Meta 5xx errors.'),
    ('Supabase',           'PostgreSQL database + Auth (JWT) + Vault (encrypted tokens) + Edge Functions (cron).'),
], col_w=(2.0,5.3), header='Integration Summary')


# ════════════════════════════════════════════════════════════
#  9. AI AGENT  (unchanged)
# ════════════════════════════════════════════════════════════
H('9. AI Agent Design  (unchanged from v1.0)', 1, C_INDIGO, pb=True)
divider()
body('No changes. Full specification in v1.0 §6. Key parameters: confidence threshold 0.75, approval state machine (proposed → approved/edited/rejected/expired), meeting parser.')


# ════════════════════════════════════════════════════════════
#  10. DAILY REPORT  (unchanged)
# ════════════════════════════════════════════════════════════
H('10. Daily WhatsApp Report  (unchanged from v1.0)', 1, C_INDIGO)
divider()
body('No changes. Full specification in v1.0 §7. Default: 07:00 Asia/Jerusalem. Sections: TODAY / NEXT 7 DAYS / AI PENDING / MEETINGS / YESTERDAY WINS.')


# ════════════════════════════════════════════════════════════
#  11. SECURITY  (unchanged)
# ════════════════════════════════════════════════════════════
H('11. Security Design  (unchanged from v1.0)', 1, C_INDIGO)
divider()
body('No changes. Full specification in v1.0 §10. Controls: JWT auth, HMAC webhook verification, Supabase Vault for OAuth tokens, WA number whitelist, audit_log, TLS 1.3.')


# ════════════════════════════════════════════════════════════
#  12. BUILD PHASES  (updated)
# ════════════════════════════════════════════════════════════
H('12. Build Phases  (v2 — UI-Aligned)', 1, C_INDIGO, pb=True)
divider()

H('Phase 1 — MVP: WhatsApp Core  (unchanged)', 2, C_CYAN)
body('Deliverable: WhatsApp bot that creates tasks (text + voice) + daily report.', bold=True)
for item in [
    'FastAPI scaffold + Supabase DB (tasks, mailboxes, scheduled_reports)',
    'WhatsApp webhook + HMAC verification',
    'Text + voice task creation pipeline',
    'Daily report cron job',
    'Seed two mailboxes (lbatech + mepsltn)',
]: bullet(item)
sp()

H('Phase 2 — REST API for Gemini UI  (expanded in v2)', 2, C_CYAN)
body('Deliverable: Full REST API including all Gemini UI endpoints. OpenAPI spec published.', bold=True)
for item in [
    'All /api/v1/tasks endpoints with Kanban envelope, timeline view, list view',
    'GET /dashboard/stats + /dashboard/recent-contacts (progress bar + avatar row)',
    'GET /tasks/{id}/messages + /tasks/{id}/activity (centre panel)',
    'PATCH /tasks/reorder (Kanban drag-drop)',
    'GET /mailboxes/{id}/inbox + /sync + /flag (right panel Outlook tab)',
    'Computed fields on every task response: due_status, due_badge, priority_color, source_icon, source_label',
    'Supabase Auth + JWT middleware + CORS for Gemini UI domain',
    'OpenAPI / Swagger docs at /docs',
]: bullet(item)
sp()

H('Phase 3 — AI Agent + Meeting Scheduler  (unchanged)', 2, C_CYAN)
body('Deliverable: Agent monitors WA messages. Meetings schedulable by voice/text.', bold=True)
for item in [
    'AI agent intent classifier + approval_queue',
    'WA approval flow: ✅  ✏️  ❌',
    'Meeting parser + meetings table',
    'MS Graph calendar create / update / cancel',
    'Multi-mailbox token refresh worker',
]: bullet(item)
sp()

H('Phase 4 — Contact Directory + Outlook Flag Sync  (unchanged)', 2, C_CYAN)
body('Deliverable: Contact directory. Flagged emails auto-create tasks. Avatar URLs on cards.', bold=True)
for item in [
    'companies + contacts tables + CSV/vCard import',
    'contacts.avatar_url + contacts.initials fields',
    'Fuzzy name matcher (AI agent + meeting parser)',
    'MS Graph subscription for mail flag → task creation',
    'Add-mailbox OAuth self-service flow',
]: bullet(item)
sp()


# ════════════════════════════════════════════════════════════
#  13. OPEN QUESTIONS  (updated)
# ════════════════════════════════════════════════════════════
H('13. Open Questions  (v2)', 1, C_INDIGO, pb=True)
divider()
tbl([
    ('[P1 — Owner]  WA number type',      'Use personal WA via Twilio Business API or Meta Cloud API dedicated number?'),
    ('[P1 — Owner]  Daily report time',   'Confirm time. Default: 07:00 Asia/Jerusalem.'),
    ('[P1 — Owner]  Bot language',        'Hebrew, English, or auto-mirror the sender language?'),
    ('[P2 — Gemini] Web UI domain',       'Confirm domain for CORS + TLS configuration.'),
    ('[P2 — Gemini] API base URL',        'Confirm API base URL so Gemini UI environment variables can be set.'),
    ('[P2 — Gemini] Avatar upload',       'Will contacts have photo upload in UI? If yes, add POST /contacts/{id}/avatar endpoint.'),
    ('[P2 — Gemini] Timeline date range', 'Default Gantt view range: 1 week, 2 weeks, or 1 month?'),
    ('[P3 — Owner]  Meeting duration',    'Default when not stated: 30 min or 60 min?'),
    ('[P4 — Owner]  Contact sample CSV',  'Provide sample export for import parser validation.'),
], col_w=(2.8,4.5), header='Open Questions')


# ════════════════════════════════════════════════════════════
#  14. GLOSSARY  (updated)
# ════════════════════════════════════════════════════════════
H('14. Glossary', 1, C_INDIGO)
divider()
tbl([
    ('TTD',          'Task Tracker Dashboard — the product name'),
    ('L1–L5',        'Architecture layers: Ingestion, AI Processing, Domain, Storage, Delivery'),
    ('AI Agent',     'Background process reading inbound WA messages and proposing tasks'),
    ('Kanban',       'Left panel of the Gemini dashboard. Three columns: To-Do, In-Progress, Done'),
    ('Gantt',        'Timeline view in right panel. Tasks rendered as date-range bars.'),
    ('due_status',   'Server-computed field: overdue | due_today | due_soon | upcoming | completed | no_date'),
    ('source_icon',  '"whatsapp" | "outlook" | "web" | "voice" — UI maps to icon component'),
    ('short_name',   'Mailbox abbreviation shown as IN badge on task cards (e.g., "lbatech")'),
    ('Approval Queue','Staging table for AI-proposed tasks awaiting owner confirmation'),
    ('Whisper',      'OpenAI speech-to-text for WA voice notes'),
    ('GPT-4o',       'OpenAI LLM for NLP, intent classification, meeting parsing'),
    ('Graph API',    'Microsoft Graph REST API for Outlook calendar + mail'),
    ('P1 / P2 / P3', 'Priority: 1=Business (red), 2=Personal (yellow), 3=General (green)'),
    ('WA',           'WhatsApp'),
], col_w=(1.5,5.8), header='Glossary')


# ── Save ─────────────────────────────────────────────────────
out = '/sessions/brave-kind-johnson/mnt/TTD/TTD_Product_Architecture_v2.0.docx'
doc.save(out)
print(f'Saved: {out}')
